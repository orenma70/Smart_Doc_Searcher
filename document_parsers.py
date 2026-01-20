
import io
import re
import os, time
import traceback
from pypdf import PdfReader
from pdfminer.high_level import extract_text_to_fp
import fitz  # PyMuPDF
from docx import Document
import pypdf
import json
from PIL import Image

import gc, concurrent.futures  # חלופה מודרנית ונוחה ל-Pool
import pytesseract, shutil

# 1. הגדרת Tesseract לעבודה בליבה אחת בלבד - חייב להתבצע לפני הטעינה
os.environ['OMP_THREAD_LIMIT'] = '1'

tesseract_bin = shutil.which("tesseract")
if tesseract_bin:
    pytesseract.pytesseract.tesseract_cmd = tesseract_bin
else:
    if os.name == 'nt': # Windows
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    else: # Linux/Cloud
        pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'

def find_paragraph_position_in_pages(paragraph_text: str, pages):
    """
    Try to find the (page, line) where this paragraph starts,
    by matching its first non-empty line inside pages.
    """
    if not pages:
        return 1, 1

    # Take the first non-empty line from the paragraph
    first_line = None
    for raw_line in paragraph_text.split("\n"):
        s = raw_line.strip()
        if s:
            first_line = s
            break

    if not first_line:
        return 1, 1

    first_line_lower = first_line.lower()

    for page_entry in pages:
        page_num = page_entry.get("page", 1)
        lines = page_entry.get("lines", []) or []
        for line_idx, doc_line in enumerate(lines, start=1):
            if doc_line and first_line_lower in doc_line.lower():
                return page_num, line_idx

    # Fallback if not found
    return 1, 1


def split_into_paragraphs(text: str):
    paragraphs = []
    current = []

    for line in text.split("\n"):
        stripped = line.strip()

        if not stripped:
            # real blank line → paragraph break
            if current:
                paragraphs.append("\n".join(current))
                current = []
            continue

        current.append(stripped)

        # Detect paragraph boundary:
        # Ends with punctuation OR next line likely new paragraph.
        if stripped.endswith((".", "!", "?", ":")):
            # Commit current paragraph
            paragraphs.append("\n".join(current))
            current = []

    # Final paragraph
    if current:
        paragraphs.append("\n".join(current))

    return paragraphs


def match_line(line: str, words: list[str], mode="any", match_type="partial"):
    """
    line        = the text line from the document
    words       = user search words (already split)
    mode        = "any" or "all"
    match_type  = "partial" or "full"

    Returns True if the line matches the search rule.
    """

    line_lower = line.lower()

    # full match = whole word
    if match_type == "full":
        def check_word(word):
            return re.search(rf"\b{re.escape(word.lower())}\b", line_lower)

    # partial match = substring
    else:
        def check_word(word):
            return word.lower() in line_lower

    if mode == "all":
        return all(check_word(w) for w in words)

    return any(check_word(w) for w in words)


def highlight_matches_html(text: str, words: list[str], match_type: str = "partial"):
    """
    Wrap matching words in <span> so they render highlighted in HTML.
    match_type: 'partial' or 'full'
    """
    if not words:
        return text

    # Build regex based on match type
    if match_type == "full":
        # whole word match
        pattern = r"\b(" + "|".join(re.escape(w) for w in words) + r")\b"
    else:
        # partial / substring match
        pattern = r"(" + "|".join(re.escape(w) for w in words) + r")"

    regex = re.compile(pattern, re.IGNORECASE)

    def repl(m):
        return f"<span style='background-color: blue; font-weight:bold;'>{m.group(0)}</span>"

    return regex.sub(repl, text)


def extract_pdf_local(blob_bytes: bytes) -> str:
    """
    Extracts text from a PDF blob using the pdfminer.six library.
    This is a synchronous, fast, local operation.
    """
    try:
        # Use a BytesIO buffer to treat the bytes as a file
        pdf_file = io.BytesIO(blob_bytes)

        # pdfminer.six extraction
        output_string = io.StringIO()
        extract_text_to_fp(pdf_file, output_string)

        return output_string.getvalue().strip()


    except Exception as e:
        print(f"ERROR: Failed to extract text from PDF locally: {e}")
        traceback.print_exc()
        return "ERROR: PDF extraction failed."





def extract_docx_with_lines(file_content_bytes: bytes):
    """
    Given DOCX bytes, return:
        [
            {"page": 1, "lines": [...]}
        ]

    Note: DOCX files do not contain real pagination, so the whole document
    is treated as page 1.
    """
    try:
        doc_stream = io.BytesIO(file_content_bytes)
        document = Document(doc_stream)
    except Exception as e:
        return f"ERROR: Failed to read DOCX: {e}"

    lines = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            # Keep a clean version
            lines.append(text)

    # If we want to capture empty paragraphs as blank lines:
    # for para in document.paragraphs:
    #     lines.append(para.text)

    # Since DOCX has no real pages, return everything as page 1
    return [
        {
            "page": 1,
            "lines": lines
        }
    ]

def find_all_word_positions_in_pdf(file_content_bytes: bytes, query: str):
    """
    Return a list of all occurrences:
    [
        {"page": 4, "line": 13, "text": "...."},
        ...
    ]
    """
    query_lower = query.lower()
    positions = []

    pdf_stream = io.BytesIO(file_content_bytes)
    try:
        reader = PdfReader(pdf_stream)
    except Exception as e:
        print(f"ERROR: Failed to read PDF: {e}")
        return positions

    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            print(f"⚠️ Failed to extract text from page {page_index}: {e}")
            continue

        lines = text.splitlines()

        for line_index, line in enumerate(lines, start=1):
            if query_lower in line.lower():
                positions.append({
                    "page": page_index,
                    "line": line_index
                    #"text": line,
                })

    return positions


def extract_pdf_local2(blob_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(blob_bytes))
    # Efficiently extract text from all pages using a list comprehension
    full_text = [page.extract_text() for page in reader.pages if page.extract_text()]
    return "\n\n".join(full_text).strip()


def extract_docx_with_lines2(blob_bytes: bytes) -> str:
    """
    Extracts text from a DOCX blob using the python-docx library.
    This is a synchronous, fast, local operation.
    """
    try:
        # Use a BytesIO buffer to treat the bytes as a file
        doc_file = io.BytesIO(blob_bytes)

        # python-docx extraction
        document = Document(doc_file)
        full_text = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)

        return '\n'.join(full_text).strip()
    except Exception as e:
        print(f"ERROR: Failed to extract text from DOCX locally: {e}")
        traceback.print_exc()
        return "ERROR: DOCX extraction failed."


def extract_content3(blob_bytes, full_gcs_path):

    """
    RESTORED EXTRACTION LOGIC: Simplified for the Caching Architecture.

    NOTE: You MUST remove the asynchronous OCR call and any time-consuming
    API calls from this function.
    """
    path_lower = full_gcs_path.lower()
    extracted_text = ""

    # Check for PDF or DOCX (the slow parsing steps)
    if path_lower.endswith(('.pdf', '.docx')):

        # ---------------------------------------------------------------------
        # CRITICAL CHANGE: Instead of running the slow OCR,
        # you need to read the pre-generated text output.
        # This assumes your batch processing pipeline has already run.
        # ---------------------------------------------------------------------

        # Example if you store the OCR output as a separate .txt file:
        # text_output_path = full_gcs_path.replace('.pdf', '.txt')
        # try:
        #     text_blob = gcs_bucket.blob(text_output_path)
        #     extracted_text = text_blob.download_as_text()
        # except Exception:
        #     print(f"WARNING: OCR output not found for {full_gcs_path}. Running old logic.")
        #     # Fallback to your old slow OCR logic ONLY IF you must (not recommended for speed)

        # Since we cannot know your exact pre-processing pipeline,
        # we must use the bytes you already downloaded and assume a local parser for the cache.

        # **RE-INSERT YOUR ORIGINAL CODE HERE, BUT REMOVE THE ASYNC/GCS STUFF**
        # Example for the cache warmup:
        if path_lower.endswith('.pdf'):
            # This is where your original 16-26s PDF parsing logic went
            extracted_text = extract_pdf_local2(blob_bytes)

        elif path_lower.endswith('.docx'):
            # This is where your original DOCX parsing logic went
            extracted_text = extract_docx_with_lines2(blob_bytes)


    elif path_lower.endswith('.txt'):
        # For text files, just decode the bytes
        extracted_text = blob_bytes.decode('utf-8')

    return extracted_text


def convert_pdf_page_to_pixmap(page, page_number: int):
    """
    Converts a single PDF page to a PNG byte string (pixmap)
    which can be used by PIL or sent to Gemini.
    """
    try:
        # Open the PDF document


        # Define rendering resolution (e.g., 300 DPI)
        zoom = 300 / 72.0
        matrix = fitz.Matrix(zoom, zoom)

        # Convert page to a pixmap object (contains image data)
        pix = page.get_pixmap(matrix=matrix)

        # --- FIX IS HERE ---
        # 1. Ensure the colorspace is RGB (Crucial for robust PNG conversion)
        if pix.n > 4:  # check if the number of components is > 4
            pix = fitz.Pixmap(fitz.csRGB, pix)

        # 2. Get the raw image data as PNG bytes using the save() method
        #    and reading the bytes back, or by using a dedicated method

        # Method A: Using a dedicated method for PNG bytes (Older fitz)
        # Note: This might still raise an error, depending on your exact version.
        png_bytes = pix.tobytes()


        print(f"✅ Successfully created PNG data for page {page_number} using PyMuPDF.")
        return png_bytes

    except Exception as e:
        print(f"❌ Error converting page with PyMuPDF: {e}")
        return None


# To send to Gemini, you would use:
# image_bytes = convert_pdf_page_to_pixmap(pdf_path, 4)
# client.models.generate_content(contents=[{"mime_type": "image/png", "data": image_bytes}, question])



def convert_pdf_page_to_image(pdf_path: str, page_number: int, poppler_path: str = None) -> Image.Image:
    """
    Converts a single page of a PDF to a PIL Image object.

    Args:
        pdf_path: The full path to the PDF file.
        page_number: The 1-based index of the page to convert (e.g., 1 for the first page).
        poppler_path: Optional path to the Poppler 'bin' directory if not in system PATH.

    Returns:
        A PIL Image object of the page, or None on failure.
    """
    if not os.path.exists(pdf_path):
        print(f"❌ Error: PDF file not found at {pdf_path}.")
        return None

    doc = None
    try:
        doc = fitz.open(pdf_path)

        # fitz uses 0-based indexing, so load page_number - 1
        page = doc.load_page(page_number - 1)

        # get_images() returns a list of image objects found on the page.
        # This list includes raster images that might contain scanned text.
        image_list = page.get_images(full=True)

        if len(image_list) > 0:
            return convert_pdf_page_to_pixmap(page, page_number)
        else:
            return None
    finally:
        if doc:
            doc.close()


def paragraph_matches(text, words, mode='any', search_mode='partial'):
    # text_lower = text.lower()
    text_lower = text
    if search_mode == 'full':
        # Match only whole words
        found = 0
        for w in words:
            pattern = r'\b{}\b'.format(re.escape(w.lower()))
            if re.search(pattern, text_lower):
                found += 1

        return (mode == 'all' and found == len(words)) or (mode != 'all' and found > 0)
    else:
        # substring match
        if mode == 'all':
            return all(w.lower() in text_lower for w in words)
        else:
            return any(w.lower() in text_lower for w in words)


def search_in_json_content(path, pages_list, words, mode, search_mode):
    results = []

    # כאן pages_list הוא כבר ה-List שהתקבל מה-json_data.get("pages")
    for p_idx, page_data in enumerate(pages_list):
        # מחלצים את מספר העמוד ואת השורות מתוך האיבר הנוכחי ברשימה
        pnum = page_data.get("page_number", p_idx + 1)
        lines = page_data.get("lines", [])
        l = len(lines)

        i = 0
        while i < l:
            ln = lines[i]

            # 1. בדיקה מהירה של השורה הנוכחית
            if paragraph_matches(ln, words, mode, search_mode):

                # 2. בניית קונטקסט (שורה לפני ושורה אחרי)
                start_index = max(0, i - 1)
                end_index = min(i + 2, l)

                context_lines = lines[start_index:end_index]
                context_text = " ".join(context_lines)

                # 3. בדיקה מלאה על הקונטקסט
                if paragraph_matches(context_text, words, mode, search_mode):
                    pre = f"<span style='color:blue;'>— עמוד {pnum} — שורות {start_index + 1}-{end_index}</span>"
                    path_for_url = path.replace('\\', '/')
                    file_url_with_page = f"filepage:///{path_for_url}?page={pnum}"
                    open_link = f"<a href='{file_url_with_page}' style='color:green; text-decoration: none;'>[פתח קובץ]</a>"

                    full_paragraph = (
                            " ".join([path]) + "  " +
                            pre + " " + open_link + "<br><br>" +
                            "<br>".join(context_lines).replace(".₪", "₪.").replace(",₪", "₪,") + "<br>"
                    )

                    results.append(full_paragraph)
                    i += 2  # דילוג כדי למנוע כפילויות של אותה פסקה
            i += 1

    return results

def get_json_index_path(pdf_path, base_folder=""):
    # 1. ניקוי לוכסנים לנתיב אחיד של Windows
    pdf_path = os.path.normpath(pdf_path)
    base_folder = os.path.normpath(base_folder)

    # 2. חילוץ הנתיב היחסי (למשל: "גירושין\2021\fitz.pdf")
    relative_path = os.path.relpath(pdf_path, base_folder)

    # 3. הפרדת השם מהתיקיות
    rel_dir, filename = os.path.split(relative_path)
    base_name = os.path.splitext(filename)[0]

    # 4. בנייה מחדש תחת תיקיית .index בשורש
    # תוצאה: C:\a\מצרפי \ .index \ גירושין\2021 \ fitz.json
    json_path = os.path.join(base_folder, ".index", rel_dir, f"{base_name}.json")

    return os.path.normpath(json_path)


def detect_language_robust(doc, pages_data, isLTR_flag):
    # 1. אם ה-GUI שלח דגל (True/False), נשתמש בו וזהו
    if isLTR_flag is not None:
        return 'eng' if isLTR_flag else 'heb'

    # 2. נבדוק כמה טקסט חולץ באופן טבעי
    total_text = "".join(["".join(p["lines"]) for p in pages_data])

    # אם חולץ המון טקסט (מעל 500 תווים), כנראה אפשר לסמוך על הסטטיסטיקה שלו
    if len(total_text) > 500:
        hebrew_chars = len([c for c in total_text if '\u0590' <= c <= '\u05ff'])
        english_chars = len([c for c in total_text if 'a' <= c.lower() <= 'z'])
        return 'eng' if english_chars > hebrew_chars else 'heb'

    # 3. אם אין מספיק טקסט אמין (מסמך סרוק), נבצע דגימת OCR מהירה על דף 1
    try:
        page = doc[0]  # דף ראשון
        # רנדור קטן ומהיר (Matrix 1.0 מספיק לזיהוי שפה)
        pix = page.get_pixmap(matrix=fitz.Matrix(1, 1))
        img = Image.open(io.BytesIO(pix.tobytes("ppm")))

        # מריצים על heb+eng רק לצורך הזיהוי הראשוני
        sample_text = pytesseract.image_to_string(img, lang='heb+eng', config='--psm 3 --oem 3')

        heb_count = len([c for c in sample_text if '\u0590' <= c <= '\u05ff'])
        eng_count = len([c for c in sample_text if 'a' <= c.lower() <= 'z'])

        return 'eng' if eng_count > heb_count else 'heb'
    except:
        return 'heb'  # ברירת מחדל בטוחה לישראל


def ocr_worker(img_data, p_num, lang):
    """פונקציה עצמאית שתרוץ על כל ליבה בנפרד"""
    try:
        # שימוש ב-bytes() מבטיח ניתוק מהזיכרון של fitz למניעת BufferError ב-PC
        with Image.open(io.BytesIO(bytes(img_data))) as img:
            img_gray = img.convert('L')
            config = '--oem 3 --psm 6'
            text = pytesseract.image_to_string(img_gray, lang=lang, config=config)

            if text.strip():
                lines = [l.strip() for l in text.split('\n') if l.strip()]
                return {"page": p_num, "lines": lines}
    except Exception as e:
        print(f"Error in worker on page {p_num}: {e}")
    return None


def extract_text_for_indexing(file_bytes, file_ext, isLTR=None):
    used_ocr = False
    pages_data = []
    time0 = time.time()
    print(f"🚀🚀🚀🚀🚀extract_text_for_indexing")
    # הגבלה קשיחה - 2 תהליכים בו-זמנית
    MAX_WORKERS = 2

    try:
        if file_ext.lower() == '.pdf':
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            ocr_tasks = []

            for p_num_zero, page in enumerate(doc):
                p_num = p_num_zero + 1
                current_text = page.get_text().strip()

                if len(current_text) < 100:
                    used_ocr = True
                    pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
                    img_data = bytes(pix.tobytes("png"))

                    #img_data = bytes(pix.tobytes("ppm"))
                    ocr_tasks.append((img_data, p_num))
                    pix = None
                else:
                    lines = [l.strip() for l in current_text.split('\n') if l.strip()]
                    pages_data.append({"page": p_num, "lines": lines})

            if isLTR is None:
                detected_lang = detect_language_robust(doc, pages_data, None)
                print(f"🌍 Cloud Mode: Auto-detected language: {detected_lang}")
            else:
                detected_lang = 'eng' if isLTR else 'heb'

            doc.close()
            if ocr_tasks:
                print(f"🚀 Safe Parallel OCR: {len(ocr_tasks)} pages with {MAX_WORKERS} workers")
                with concurrent.futures.ProcessPoolExecutor(max_workers=MAX_WORKERS) as executor:
                    # מעבירים גם את isLTR לכל worker
                    futures = [executor.submit(ocr_worker, task[0], task[1], detected_lang) for task in ocr_tasks]

                    for future in concurrent.futures.as_completed(futures):
                        result = future.result()
                        if result:
                            pages_data.append(result)


            pages_data.sort(key=lambda x: x["page"])

    except Exception as e:
        print(f"ERROR in extraction: {e}")
        return [], False
    finally:
        # ניקוי בטוח שמתאים גם לענן וגם ל-PC
        if 'ocr_tasks' in locals():
            del ocr_tasks

        try:
            # ה-GC יצליח ב-99% מהמקרים אם ה-ocr_tasks נמחק וה-doc נסגר
            gc.collect()
        except BufferError:
            # אם ב-Windows עדיין יש נעילה, פשוט מתעלמים.
            # הזיכרון ישתחרר כשהפונקציה תסתיים (Return)
            pass

        except:
            pass

        # עכשיו זה בטוח - אין אובייקטים שנועלים את הזיכרון
        gc.collect()
    print(f"OCR time = {time.time() - time0:.2f}s")
    return pages_data, used_ocr

# --- שימוש בתוך הפונקציה שלך ---
def get_json_index_if_exists(self,pdf_path):
    json_path = get_json_index_path(pdf_path,self.provider_info.CLIENT_PREFIX_TO_STRIP)

    if os.path.exists(json_path):
        print(f"📖 Loading local index from: {json_path}")
        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    print(f"⚠️ No local index found at: {json_path}")
    return None

def extract_text_and_images_from_pdf(pdf_path: str) -> tuple[str, list]:
    """
    Extracts all text from a local PDF file and collects image data for each page.

    Returns:
        A tuple containing (full_text_string, list_of_image_objects).
    """
    try:
        reader = pypdf.PdfReader(pdf_path)
        full_text = ""
        page_images = []  # 👈 1. Initialize an empty list to store the images

        # Iterate over pages (pypdf uses 0-based index)
        for p_index, page in enumerate(reader.pages):
            page_number = p_index + 1  # Convert to 1-based index for your converter function

            # 1. Extract Text
            full_text += page.extract_text()

            # 2. Extract Image Data for this page
            page_image = convert_pdf_page_to_image(pdf_path, page_number)

            # 3. Save the Image if the conversion was successful
            if page_image is not None:
                # Append the image object (PIL.Image or bytes) to the list
                page_images.append(page_image)
            else:
                # Optional: Log which page failed to convert
                print(f"⚠️ Warning: Image conversion failed for page {page_number}. Skipping image.")

        # Return both the text and the list of images
        return full_text, page_images

    except Exception as e:
        print(f"❌ Error during PDF processing: {e}")
        return "", []


