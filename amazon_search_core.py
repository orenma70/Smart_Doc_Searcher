import boto3, os
import time
import traceback
from flask import Flask, request, jsonify
import io
import fitz  # PyMuPDF - כבר נמצא ב-requirements שלך
from docx import Document  # כבר נמצא ב-requirements שלך
import pytesseract
from amazon_search_utilities import match_line, highlight_matches_html, search_in_json_content
from PIL import Image
import json
import config_reader
from document_parsers import extract_text_for_indexing

cloud_provider="Amazon"
PROVIDER_CONFIG=config_reader.set_provider_config(cloud_provider)

BUCKET_NAME = PROVIDER_CONFIG["BUCKET_NAME"]






app = Flask(__name__)
s3 = boto3.client('s3')
# --- AWS Configuration ---

def get_documents_for_path(directory_path):
    documents = []
    try:
        paginator = s3.get_paginator('list_objects_v2')
        base_prefix = directory_path.strip('/') + '/' if directory_path else ""

        for page in paginator.paginate(Bucket=BUCKET_NAME, Prefix=base_prefix):
            if 'Contents' not in page: continue

            for obj in page['Contents']:
                key = obj['Key']
                filename = key.split('/')[-1]

                # Skip folders and existing index folder
                if key.endswith('/') or filename.startswith('~$') or key.startswith('.index/'):
                    continue

                base_path = key.rsplit('.', 1)[0] if '.' in key else key
                index_key = f".index/{base_path}.json".replace("//", "/")

                pages = []
                try:
                    # 1. Attempt to load JSON Index
                    idx_resp = s3.get_object(Bucket=BUCKET_NAME, Key=index_key)
                    index_data = json.loads(idx_resp['Body'].read().decode('utf-8'))
                    raw_pages = index_data.get("pages", [])

                    # 2. FIX: Ensure pages is a list of DICTS, not strings (Fixes AttributeError)
                    for idx, p in enumerate(raw_pages):
                        if isinstance(p, str):
                            pages.append({"page_number": idx + 1, "lines": [p]})
                        else:
                            pages.append(p)



                except s3.exceptions.NoSuchKey:

                    file_obj = s3.get_object(Bucket=BUCKET_NAME, Key=key)

                    file_content = file_obj['Body'].read()

                    file_ext = filename.lower()

                    pages = []  # אתחול תמיד

                    print(f"🔍 Index missing for {filename}. Extracting real content...")

                    if file_ext.endswith('.docx'):

                        doc_reader = Document(io.BytesIO(file_content))

                        extracted_text = "\n".join([para.text for para in doc_reader.paragraphs])

                        pages = [{"page_number": 1, "lines": extracted_text.splitlines()}]


                    elif file_ext.endswith('.pdf'):

                        with fitz.open(stream=file_content, filetype="pdf") as pdf:

                            num_pages = len(pdf)

                            # חילוץ טקסט דיגיטלי ראשוני

                            full_digital_text = "\n".join([p.get_text() for p in pdf])

                        # בדיקת סף OCR

                        avg_chars = len(full_digital_text) / max(num_pages, 1)

                        if avg_chars < 200:

                            print(f"🚀 OCR triggered (Avg chars: {avg_chars:.1f})")

                            raw_pages, was_ocr = extract_text_for_indexing(file_content, '.pdf', isLTR=None)
                            pages = [{"page_number": p["page"], "lines": p["lines"]} for p in raw_pages]

                            index_data = {
                                "filename": filename,
                                "pages": raw_pages,
                                "timestamp": time.time()
                            }
                            base_name = os.path.splitext(filename)[0]
                            clean_prefix = base_prefix.strip("/")
                            local_index_path = os.path.join(".index", clean_prefix, f"{base_name}.json")

                            try:
                                target_dir = os.path.dirname(local_index_path)
                                if not os.path.exists(target_dir):
                                    print(f"📂 DEBUG: Creating missing directory: {target_dir}")
                                    os.makedirs(target_dir, exist_ok=True)

                                json_payload = json.dumps(index_data, ensure_ascii=False, indent=4).encode('utf-8')

                                s3.put_object(
                                    Bucket=BUCKET_NAME,
                                    Key=index_key,  # הנתיב שמתחיל ב-.index/
                                    Body=json_payload,
                                    ContentType='application/json'
                                )



                            except Exception as save_error:
                                print(f"🔥 DEBUG: Failed to write local index: {save_error}")


                        else:

                            with fitz.open(stream=file_content, filetype="pdf") as pdf:

                                for i, page in enumerate(pdf):
                                    # כאן page.get_text() הוא string, אז splitlines עובד

                                    pages.append({

                                        "page_number": i + 1,

                                        "lines": page.get_text().splitlines()

                                    })


                    else:

                        try:

                            extracted_text = file_content.decode('utf-8')

                        except:

                            extracted_text = file_content.decode('cp1255', errors='ignore')

                        pages = [{"page_number": 1, "lines": extracted_text.splitlines()}]

                except Exception as e:
                    print(f"⚠️ Error parsing index for {filename}: {e}")
                    continue

                # 4. Success: Document is now searchable even without a perfect index
                documents.append({
                    "name": filename,
                    "full_path": key,
                    "pages": pages
                })
        return documents
    except Exception as e:
        print(f"🔥 S3 Error: {str(e)}")
        return []



def simple_keyword_search(query, directory_path="", mode="any", match_type="partial", show_mode="paragraph"):
    documents = get_documents_for_path(directory_path)
    if not documents:
        return {"status": "ok", "details": "No documents found", "matches": []}

    words = [w.strip() for w in query.split() if w.strip()]
    results = []

    for doc in documents:
        # אם אנחנו ב-Paragraph Mode, נשתמש בלוגיקה של ה-GUI
        if show_mode == "paragraph":
            matches_html = search_in_json_content(
                doc["full_path"], doc.get("pages", []), words, mode, match_type
            )
            if matches_html:
                results.append({
                    "file": doc["name"],
                    "full_path": doc["full_path"],
                    "matches_html": matches_html,
                    "match_positions": []
                })
        else:  # Line Mode
            matched_items_html = []
            for page_entry in doc.get("pages", []):
                p_num = page_entry.get("page", 1)
                for line in page_entry.get("lines", []):
                    if match_line(line, words, mode, match_type):
                        # הוספת מספר העמוד לכל שורה שנמצאה
                        highlighted = highlight_matches_html(line, words, match_type)
                        matched_items_html.append(f"עמוד {p_num}: {highlighted}")

            if matched_items_html:
                results.append({
                    "file": doc["name"],
                    "full_path": doc["full_path"],
                    "matches_html": matched_items_html
                })

    return {"status": "ok", "query": query, "matches": results}


@app.route('/')
def home():
    return {"status": "searching_api_active"}, 200
# ---------------------

@app.route('/simple_search', methods=['POST'])
def simple_search_endpoint():
    timer_start = time.time()
    data = request.get_json(silent=True) or {}

    query = data.get('query', '').strip()
    directory_path = data.get('directory_path', '').strip()
    config = data.get("search_config", {})

    try:
        result = simple_keyword_search(
            query, directory_path,
            mode=config.get("word_logic", "any"),
            match_type=config.get("match_type", "partial"),
            show_mode=config.get("show_mode", "paragraph")
        )
        result["debug"] = f"{round(time.time() - timer_start, 2)} sec"
        return jsonify(result), 200
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/version')
def get_version():
    # הוא ימשוך אוטומטית את ה-v21.1.0 מה-App Runner
    return {
        "version": os.getenv("APP_VERSION", "local-dev"),
        "status": "stable",
        "mode": "paragraph"
    }

if __name__ == "__main__":
    if False:
        app.run(host='0.0.0.0', port=8080)
    else:
        #test_docx_parsing("C:\\a\\b.docx")
        file_path = "C:\\a\\b.docx"
        filename = "C:\\a\\b.pdf"
        with open(file_path, 'rb') as f:
            file_bytes = f.read()

        with open(filename, "rb") as f:
            pdf_bytes = f.read()

        file_ext = os.path.splitext(filename)[1].lower()
        pages_data, was_ocr_needed = extract_text_for_indexing(pdf_bytes, file_ext)


        doc_obj = Document(io.BytesIO(file_bytes))
        all_lines = []
        # יצירת רשימת עמודים אמיתית יותר ל-GUI
        final_pages = []

        # 1. טקסט רגיל (עמודים ראשונים)
        text_lines = []
        for p in doc_obj.paragraphs:
            if p.text.strip():
                text_lines.append(p.text.strip())

        all_lines.extend(text_lines)
        final_pages.append({"page": 1, "lines": text_lines})

        # 2. OCR על תמונות (עמוד 4 ואילך)
        image_counter = 2  # נתייחס לכל תמונה כעמוד חדש לצורך התצוגה
        for rel_id, rel in doc_obj.part.rels.items():
            target_ref_safe = getattr(rel, 'target_ref', '')
            if "image" in target_ref_safe or "/media/" in target_ref_safe:
                try:
                    img = Image.open(io.BytesIO(rel.target_part.blob)).convert("L")
                    ocr_text = pytesseract.image_to_string(img, lang='heb')

                    if ocr_text.strip():
                        lines = [l.strip() for l in ocr_text.split('\n') if l.strip()]
                        all_lines.extend(lines)
                        # כאן התיקון! מוסיפים עמוד חדש ל-GUI
                        final_pages.append({"page": image_counter, "lines": lines})
                        image_counter += 1
                        print(f"DEBUG: OCR found text in {rel_id}")
                except:
                    continue

        final_content = "\n\n".join(all_lines)
        a=0