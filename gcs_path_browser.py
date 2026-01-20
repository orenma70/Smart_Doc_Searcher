import os, sys
from google.cloud import storage
from typing import List, Dict, Any, Optional, Union
# We assume PyQt5 is used based on QFileDialog in browse_directory
from PyQt5 import QtWidgets, QtCore, QtGui
from search_utilities import get_storage_client
from azure_search_utilities import browse_azure_path_logic
import time
import hashlib
import binascii, base64
from ui_setup import isLTR
from PyQt5.QtWidgets import QMessageBox
import boto3  # Make sure to pip install boto3
import json

from ui_setup import non_sync_cloud_str, sync_cloud_str
from azure.storage.blob import BlobServiceClient, ContentSettings
import base64
from config_reader import PROVIDER_CONFIG
from document_parsers import extract_text_for_indexing # Set your flag here based on your environment



# Global client variables
gcs_client: Optional[storage.Client] = None
s3_client: Any = None
KMS_KEY_ARN = "arn:aws:kms:ap-southeast-2:038715112888:key/82ae7f3a-eb41-4f29-bd2c-85b9ab573023"

def  get_gcs_files_recursive(bucket_name, prefix = ""):
    global gcs_client
    GCS_OCR_OUTPUT_PATH = PROVIDER_CONFIG.get("GCS_OCR_OUTPUT_PATH")

    skip_folder_name = GCS_OCR_OUTPUT_PATH.rstrip('/').split('/')[-1]

    cloud_files = {}
    if gcs_client is None: gcs_client = get_storage_client()
    bucket = gcs_client.bucket(bucket_name)
    for blob in bucket.list_blobs(prefix=prefix):
        if blob.name.startswith(f"{skip_folder_name}/") or blob.name == skip_folder_name:
            continue

        name = blob.name[len(prefix):].lstrip("/")
        if name:
            # 3. פתרון ה-MD5: המרה מ-Base64 ל-Hex
            # זה יהפוך את 'anKb...' ל-'6a72...' ויתאים ל-Local שלך
            raw_md5 = blob.md5_hash
            if raw_md5:
                cloud_files[name] = base64.b64decode(raw_md5).hex()
            else:
                cloud_files[name] = None

    return cloud_files

def get_azure_files_recursive(container_name, prefix):
    # 1. התחברות ללקוח
    connection_string = os.getenv("azuresmartsearch3key1conn")
    if not connection_string:
        print("❌ Azure Error: Connection string not found.")
        return {}

    from azure.storage.blob import BlobServiceClient
    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    container_client = blob_service_client.get_container_client(container_name)

    cloud_files = {}

    # 2. הכנת ה-search_prefix
    # אם prefix ריק או None, name_starts_with יהיה None ו-Azure יסרוק את כל ה-Container
    search_prefix = prefix if prefix and prefix.strip() else None

    # מוודא שה-prefix מסתיים ב-/ כדי לסרוק תוכן תיקייה בלבד
    if search_prefix and not search_prefix.endswith('/'):
        search_prefix += '/'

    try:
        # 3. סריקת ה-Blobs (Azure עושה Pagination אוטומטי)
        blobs = container_client.list_blobs(name_starts_with=search_prefix, include=['metadata'])

        for blob in blobs:
            full_key = blob.name

            # דילוג על התיקייה עצמה
            if search_prefix and full_key == search_prefix:
                continue

            # סינון קבצים זמניים
            file_name_only = full_key.split('/')[-1]
            if any(x in file_name_only for x in ['$~', '~$']) or file_name_only.startswith('$'):
                continue

            try:
                raw_md5 = blob.metadata.get('md5_hash') or blob.content_settings.content_md5

                if isinstance(raw_md5, bytes):
                    md5_hash = raw_md5.hex()
                elif isinstance(raw_md5, str) and '==' in raw_md5:
                    md5_hash = base64.b64decode(raw_md5).hex()
                else:
                    md5_hash = str(raw_md5).strip().lower() if raw_md5 else None

                cloud_files[full_key] = md5_hash

            except Exception as e:
                print(f"Error processing blob {full_key}: {e}")

    except Exception as e:
        print(f"Error scanning Azure container {container_name}: {e}")

    return cloud_files


def delete_from_cloud_with_index(self, filename, prefix="", skip_index_delete=True):
    """
    מוחק מהענן את הקובץ המקורי.
    מוחק את האינדקס רק אם skip_index_delete=False.
    """
    target_key = f"{prefix}/{filename}".replace("//", "/").strip("/")
    base_name = os.path.splitext(filename)[0]
    index_key = f".index/{prefix}/{base_name}.json".replace("//", "/").strip("/")

    cloud_provider = self.provider_info.get("cloud_provider")
    bucket_name = self.provider_info.get("BUCKET_NAME")

    try:
        # --- Amazon AWS ---
        if cloud_provider == "Amazon":
            client = get_s3_client()
            delete_list = [{'Key': target_key}]

            # הוספת האינדקס לרשימת המחיקה רק אם לא ביקשנו לדלג
            if not skip_index_delete:
                delete_list.append({'Key': index_key})

            client.delete_objects(
                Bucket=bucket_name,
                Delete={'Objects': delete_list, 'Quiet': True}
            )
            print(f"🗑️ AWS: Deleted {target_key} " + ("(Index skipped)" if skip_index_delete else "and its index."))

        # --- Google Cloud ---
        elif cloud_provider == "Google":
            global gcs_client
            if gcs_client is None: gcs_client = get_storage_client()
            bucket = gcs_client.bucket(bucket_name)

            # מחיקת קובץ מקורי
            blob = bucket.blob(target_key)
            if blob.exists(): blob.delete()

            # מחיקת אינדקס בתנאי
            if not skip_index_delete:
                index_blob = bucket.blob(index_key)
                if index_blob.exists(): index_blob.delete()

            print(f"🗑️ GCS: Deleted {target_key} " + ("(Index skipped)" if skip_index_delete else "and its index."))

        # --- Microsoft Azure ---
        elif cloud_provider == "Microsoft":
            from azure.storage.blob import BlobServiceClient
            connection_string = os.getenv("azuresmartsearch3key1conn")
            blob_service_client = BlobServiceClient.from_connection_string(connection_string)
            container_client = blob_service_client.get_container_client(bucket_name)

            # מחיקת קובץ מקורי
            blob_file = container_client.get_blob_client(target_key)
            if blob_file.exists(): blob_file.delete_blob()

            # מחיקת אינדקס בתנאי
            if not skip_index_delete:
                blob_index = container_client.get_blob_client(index_key)
                if blob_index.exists(): blob_index.delete_blob()

            print(f"🗑️ Azure: Deleted {target_key} " + ("(Index skipped)" if skip_index_delete else "and its index."))

        return True

    except Exception as e:
        print(f"❌ Error during cloud deletion of {filename}: {e}")
        return False


def save_json_file(self, local_folder, filename, base_folder):
    # Create the exact same path structure in the cloud
    relative_dir = os.path.relpath(local_folder, base_folder).replace("\\", "/")

    base_name = os.path.splitext(filename)[0]
    local_index_path = os.path.join(base_folder, ".index", relative_dir, f"{base_name}.json")

    os.makedirs(os.path.dirname(local_index_path), exist_ok=True)

    with open(os.path.join(local_folder, filename), "rb") as f:
        pdf_bytes = f.read()

    file_ext = os.path.splitext(filename)[1].lower()
    pages_data, was_ocr_needed = extract_text_for_indexing(pdf_bytes, file_ext)

    index_data = {
        "filename": filename,
        "pages": pages_data,
        "timestamp": time.time()
    }

    # יצירת ה-JSON כ-Bytes (למניעת בעיות Encoding/Newline בווינדוס)
    json_payload = json.dumps(index_data, ensure_ascii=False, indent=4).encode('utf-8')


    # שמירה כבינארי (wb)
    with open(local_index_path, "wb") as f:
        f.write(json_payload)



def upload_to_cloud(self, local_folder, filename, base_folder):
    # 1. Identify Cloud Provider
    cloud_provider = self.provider_info.get("cloud_provider")
    use_mode_aws = (cloud_provider == "Amazon")
    use_mode_azr = (cloud_provider == "Microsoft")
    use_mode_clr = (cloud_provider == "Google")

    # 2. Normalize paths and define target
    local_folder = os.path.normpath(local_folder)
    base_folder = os.path.normpath(base_folder)
    file_path = os.path.join(local_folder, filename)

    # Create the exact same path structure in the cloud
    relative_dir = os.path.relpath(local_folder, base_folder).replace("\\", "/")
    if relative_dir == ".": relative_dir = ""
    relative_file_path = f"{relative_dir}/{filename}".replace("//", "/").strip("/")

    base_name = os.path.splitext(filename)[0]
    local_index_path = os.path.join(base_folder, ".index", relative_dir, f"{base_name}.json")
    os.makedirs(os.path.dirname(local_index_path), exist_ok=True)

    try:
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return

        # 3. Get Hashes (Required for metadata and integrity)
        with open(file_path, "rb") as f:
            file_bytes = f.read()

        pdf_hex_md5 = hashlib.md5(file_bytes).hexdigest()
        # Using your existing helper for base64 hash (required by Azure/Google)
        hex_md5, b64_md5 = get_local_hashes(file_path)

        bucket_name = self.provider_info["BUCKET_NAME"]

        # 4. PERFORM UPLOAD
        # --- Amazon AWS ---
        if use_mode_aws:
            client = get_s3_client()
            client.upload_file(file_path, bucket_name, relative_file_path, ExtraArgs={
                'Metadata': {'md5-hash': pdf_hex_md5}
            })
            print(f"✅ Uploaded to Amazon: {filename}")

        # --- Google Cloud ---
        elif use_mode_clr:
            client = get_storage_client()
            bucket = client.bucket(bucket_name)
            blob_file = bucket.blob(relative_file_path)
            blob_file.md5_hash = b64_md5
            blob_file.metadata = {'md5-hash': pdf_hex_md5}
            blob_file.upload_from_filename(file_path)
            print(f"✅ Uploaded to Google: {filename}")

        # --- Microsoft Azure ---
        elif use_mode_azr:
            connection_string = os.getenv("azuresmartsearch3key1conn")
            if not connection_string: raise Exception("Azure connection string missing")

            from azure.storage.blob import BlobServiceClient, ContentSettings
            import base64

            blob_service_client = BlobServiceClient.from_connection_string(connection_string)
            container_client = blob_service_client.get_container_client(bucket_name)
            blob_file = container_client.get_blob_client(relative_file_path)

            md5_bytes = base64.b64decode(b64_md5)
            blob_file.upload_blob(
                file_bytes,
                overwrite=True,
                content_settings=ContentSettings(content_md5=md5_bytes),
                metadata={'md5_hash': pdf_hex_md5}
            )
            print(f"✅ Uploaded to Azure: {filename}")

    except Exception as e:
        print(f"❌ Error in upload: {e}")


def download_from_cloud(self, local_folder, filename):
    # 1. זיהוי ספק הענן
    cloud_provider = self.provider_info.get("cloud_provider")
    use_mode_aws = (cloud_provider == "Amazon")
    use_mode_azr = (cloud_provider == "Microsoft")
    use_mode_clr = (cloud_provider == "Google")

    # 2. הגדרת נתיבי יעד מקומיים
    # filename יכול להיות למשל ".index/doc1.json"
    full_local_path = os.path.join(local_folder, filename)
    bucket_name = self.provider_info["BUCKET_NAME"]

    # וודא שתיקיית היעד קיימת (למשל יצירת תיקיית .index אם היא חסרה)
    os.makedirs(os.path.dirname(full_local_path), exist_ok=True)

    try:
        # 3. ביצוע הורדה
        # --- Amazon AWS ---
        if use_mode_aws:
            client = get_s3_client()
            client.download_file(bucket_name, filename, full_local_path)
            print(f"✅ Downloaded from Amazon: {filename}")

        # --- Google Cloud ---
        elif use_mode_clr:
            client = get_storage_client()
            bucket = client.bucket(bucket_name)
            blob = bucket.blob(filename)
            blob.download_to_filename(full_local_path)
            print(f"✅ Downloaded from Google: {filename}")

        # --- Microsoft Azure ---
        elif use_mode_azr:
            connection_string = os.getenv("azuresmartsearch3key1conn")
            if not connection_string: raise Exception("Azure connection string missing")

            from azure.storage.blob import BlobServiceClient
            blob_service_client = BlobServiceClient.from_connection_string(connection_string)
            blob_client = blob_service_client.get_blob_client(container=bucket_name, blob=filename)

            with open(full_local_path, "wb") as download_file:
                download_file.write(blob_client.download_blob().readall())
            print(f"✅ Downloaded from Azure: {filename}")

    except Exception as e:
        print(f"❌ Error in download: {e}")

def get_s3_client():
    """Initializes AWS S3 client using environment variables."""
    global s3_client
    if s3_client is None:
        s3_client = boto3.client(
            's3',
            aws_access_key_id=os.environ.get("AWS_ACCESS_KEY"),
            aws_secret_access_key=os.environ.get("AWS_SECRET_KEY"),
            region_name="ap-southeast-2"
        )
    return s3_client


# ==============================================================================
# UNIFIED CLOUD BROWSER
# ==============================================================================

def browse_cloud_path(self) -> Dict[str, List[str]]:
    """Entry point for the GUI to fetch folder lists."""


    use_mode_aws = self.provider == "Amazon"
    use_mode_azr = self.provider == "Microsoft"
    use_mode_clr = self.provider == "Google"

    if use_mode_aws:
        return browse_s3_path_logic(self)
    elif use_mode_clr:
        return browse_gcs_path(self)
    elif use_mode_azr:
        return browse_azure_path_logic(self)
    else:
        return False

def browse_s3_path_logic(self) -> Dict[str, List[str]]:
    """AWS S3 implementation: returns a list of folder names just like GCS."""

    path_prefix = self.current_path
    # Example usage
    path_prefix = f"{path_prefix}"
    # 1. Normalize the prefix path
    normalized_prefix = path_prefix.strip('/').replace('\\', '/')
    if normalized_prefix and not normalized_prefix.endswith('/'):
        normalized_prefix += '/'

    prefix = normalized_prefix

    client = get_s3_client()
    try:
        # Delimiter='/' tells S3 to group files into virtual folders
        response = client.list_objects_v2(
            Bucket=self.bucket_name,
            Prefix=prefix,
            Delimiter='/'
        )

        folders = []
        # In S3, 'CommonPrefixes' are what we call folders
        if 'CommonPrefixes' in response:
            for cp in response['CommonPrefixes']:
                # Example: 'photos/vacation/' -> strip last '/' -> split -> get 'vacation'
                full_path = cp['Prefix'].rstrip('/')
                folder_name = full_path.split('/')[-1]
                folders.append(folder_name)

        return {"folders": sorted(folders)}
    except Exception as e:
        print(f"S3 Browse Error: {e}")
        return {"folders": []}


def browse_gcs_path_logic(self,prefix: str) -> Dict[str, List[str]]:
    """Your original GCS logic wrapped for the switchboard."""
    global gcs_client
    if gcs_client is None:
        gcs_client = get_storage_client()

    try:
        bucket = gcs_client.bucket(self.bucket_name)
        blobs_iterator = bucket.list_blobs(prefix=prefix, delimiter='/')
        # GCS returns folders in the 'prefixes' attribute when delimiter is used
        folders = [p.rstrip('/').split('/')[-1] for p in blobs_iterator.prefixes]
        return {"folders": sorted(folders)}
    except Exception as e:
        print(f"GCS Browse Error: {e}")
        return {"folders": []}


# ==============================================================================
# UNIFIED SYNC & UPLOAD
# ==============================================================================
def get_local_md5(file_path):
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest() # זה יחזיר 'ea49...'

def get_local_hashes(file_path):
    """Generates both Hex (S3) and Base64 (GCS) MD5 hashes."""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    hex_digest = hash_md5.hexdigest()
    b64_digest = base64.b64encode(binascii.unhexlify(hex_digest)).decode("utf-8")
    return hex_digest, b64_digest

def browse_s3_logic(self,prefix: str) -> Dict[str, List[str]]:
    """AWS S3 implementation of folder browsing."""
    try:
        # Note: Boto3 uses 'CommonPrefixes' for virtual folders
        response = s3_client.list_objects_v2(
            Bucket=self.bucket_name,
            Prefix=prefix,
            Delimiter='/'
        )
        folders = []
        if 'CommonPrefixes' in response:
            for cp in response['CommonPrefixes']:
                # Strip the prefix to get just the folder name for the UI
                full_path = cp['Prefix'].rstrip('/')
                folder_name = full_path.split('/')[-1]
                folders.append(folder_name)
        return {"folders": sorted(folders)}
    except Exception as e:
        print(f"S3 Error: {e}")
        return {"folders": []}


def get_file_hash(path):
    """Returns (Hex_Hash, Base64_Hash) to support both clouds."""
    hash_md5 = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)

    hex_md5 = hash_md5.hexdigest()
    base64_md5 = base64.b64encode(binascii.unhexlify(hex_md5)).decode("utf-8")
    return hex_md5, base64_md5


def browse_s3_path(self,prefix: str) -> Dict[str, List[str]]:
    """AWS S3 implementation of folder browsing."""
    try:
        s3_client = boto3.client('s3')  # Uses your aws configure credentials
        # Delimiter='/' is what tells S3 to 'act' like a folder system
        response = s3_client.list_objects_v2(
            Bucket=self.bucket_name,
            Prefix=prefix,
            Delimiter='/'
        )

        folders = []
        # S3 returns 'folders' in the CommonPrefixes key
        if 'CommonPrefixes' in response:
            for cp in response['CommonPrefixes']:
                # S3 returns full path 'folder/sub/', we strip to get just 'sub'
                full_name = cp['Prefix'].rstrip('/')
                folder_name = full_name.split('/')[-1]
                folders.append(folder_name)

        return {"folders": sorted(folders)}
    except Exception as e:
        print(f"S3 Error: {e}")
        return {"folders": []}



def upload_to_bucket(bucket, local_folder, filename, prefix=""):
    local_path = os.path.join(local_folder, filename)
    if prefix and not prefix.endswith('/'):
        gcs_path = f"{prefix}/{filename}"
    else:
        gcs_path = f"{prefix}{filename}"

    blob = bucket.blob(gcs_path)

    # 1. Disable "smart" content-type guessing
    # 2. Upload as a raw stream to ensure MD5 matches the HD perfectly
    blob.content_type = 'application/octet-stream'

    #with open(local_path, "rb") as f:
    blob.upload_from_filename(local_path)

def md5_of_file(path):
    """Compute MD5 checksum of a local file."""
    hash_md5 = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return base64.b64encode(binascii.unhexlify(hash_md5.hexdigest())).decode("utf-8")

def update_gcs_radio(self):
    if self.sync0:
        self.cloud_gemini_radio.setText(sync_cloud_str)
        self.display_root.setStyleSheet("color: white; background-color: black;")
    else:
        self.cloud_gemini_radio.setText(non_sync_cloud_str)
        self.display_root.setStyleSheet("color: red; background-color: black;")


import docx
import pdfplumber


def analyze_file_ocr_needs(self, full_path):
    ext = full_path.lower().split('.')[-1]

    # --- טיפול ב-PDF ---
    if ext == 'pdf':
        try:
            with pdfplumber.open(full_path) as pdf:
                total_pages = len(pdf.pages)
                scanned_pages = 0
                for page in pdf.pages:
                    # אם אין טקסט בכלל בדף, נחשיב אותו כסרוק
                    if not page.extract_text(layout=False).strip():
                        scanned_pages += 1

                if scanned_pages == total_pages: return "full_ocr"
                if scanned_pages > 0: return "partial_ocr"
                return "none"
        except:
            return "full_ocr"

    # --- טיפול ב-DOCX ---
    if ext == 'docx':
        try:
            doc = docx.Document(full_path)
            # בודקים אם יש טקסט בכלל במסמך
            full_text = "".join([p.text for p in doc.paragraphs]).strip()

            # חיפוש אלמנטים גרפיים (תמונות/סריקות)
            has_images = False
            # בדיקה בתוך פסקאות
            for paragraph in doc.paragraphs:
                if 'w:drawing' in paragraph._p.xml or 'w:pict' in paragraph._p.xml:
                    has_images = True
                    break

            # בדיקה בתוך טבלאות
            if not has_images:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if 'w:drawing' in cell._element.xml:
                                has_images = True
                                break

            if not full_text and has_images: return "full_ocr"  # מסמך ריק עם תמונה (כנראה סריקה)
            if full_text and has_images: return "partial_ocr"  # יש טקסט אבל יש גם תמונות/סריקות
            if not full_text and not has_images: return "none"  # מסמך ריק לגמרי
            return "none"
        except:
            return "partial_ocr"

    return "none"

def check_sync(self, prefix=""):

    local_path = self.provider_info.get("CLIENT_PREFIX_TO_STRIP")
    bucket_name = self.provider_info.get("BUCKET_NAME")

    # 1. איסוף קבצים מקומיים וה-Hashes שלהם (ללא שינוי)
    cloud_provider = self.provider_info["cloud_provider"]
    use_mode_aws = cloud_provider  == "Amazon"
    use_mode_azr = cloud_provider == "Microsoft"
    use_mode_clr = cloud_provider == "Google"

    local_files = {}
    for root, _, files in os.walk(local_path):
        for f in files:
            if f.startswith("$") or f.startswith("~$"): continue
            full_path = os.path.join(root, f)
            rel_path = os.path.relpath(full_path, local_path).replace("\\", "/")
            hex_md5, b64_md5 = get_local_hashes(full_path)
            md5 = get_local_md5(full_path)
            local_files[rel_path] = md5 #hex_md5 if use_mode_aws else b64_md5

    # 2. איסוף קבצים מהענן (ללא שינוי)
    cloud_files = {}
    if use_mode_aws:
        cloud_files = get_aws_cloud_files_recursive(bucket_name, prefix)
    elif use_mode_azr:
        cloud_files = get_azure_files_recursive(bucket_name, prefix)
    elif use_mode_clr:
        cloud_files = get_gcs_files_recursive(bucket_name, prefix)


    # 3. השוואת סטים (הלוגיקה המקורית שלך)
    missing_in_cloud = set(local_files) - set(cloud_files)
    missing_locally = set(cloud_files) - set(local_files)

    # רשימות לאיסוף הקבצים במקום שאלות מיידיות
    files_to_upload = []
    mismatched_files = []

    # לוגיקה לזיהוי קבצים להעלאה/עדכון
    for filename in local_files:
        is_missing = filename in missing_in_cloud
        is_json = filename.lower().endswith('.json')
        is_mismatched = False
        if not is_json:
            is_mismatched = (filename in cloud_files and local_files[filename] != cloud_files[filename])

        if is_missing or is_mismatched:
            files_to_upload.append(filename)
            if is_mismatched:
                mismatched_files.append(filename)

    if files_to_upload:
        for filename in files_to_upload:
            if not filename.lower().endswith('.json'):
                is_ocr_needed = analyze_file_ocr_needs(self, filename)
                if is_ocr_needed in {"full_ocr", "partial_ocr"}:
                    save_json_file(self,local_path, filename, base_folder=self.provider_info["CLIENT_PREFIX_TO_STRIP"])

    # --- שאלה אחת וביצוע העלאה ---
    if files_to_upload:
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Question)
        msg.setWindowTitle("סנכרון קבצים" if not isLTR else "Sync Files")
        msg.setText(
            f"נמצאו {len(files_to_upload)} קבצים להעלאה/עדכון. לבצע?" if not isLTR else f"Update {len(files_to_upload)} files?")
        msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
        if msg.exec_() == QMessageBox.Ok:
            for filename in files_to_upload:
                upload_to_cloud(self,local_path, filename, base_folder=self.provider_info["CLIENT_PREFIX_TO_STRIP"])

                if filename in missing_in_cloud:
                    missing_in_cloud.remove(filename)  # הקובץ כבר לא חסר בענן

                if filename in mismatched_files:
                    mismatched_files.remove(filename)  # הקו
    # --- שאלה אחת וביצוע מחיקה ---
    files_to_download = [f for f in missing_locally if  f.startswith(".index/")]
    if files_to_download:
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Warning)
        msg.setWindowTitle("הורדה מהענן" if not isLTR else "Download Delete")
        msg.setText(f"נמצאו {len(files_to_download)} קבצים להורדה מהענן. לבצע?" if not isLTR else f"Download {len(files_to_download)} files from cloud?")
        msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
        if msg.exec_() == QMessageBox.Ok:
            for filename in files_to_download:
                download_from_cloud(self, local_path, filename)

                if filename in missing_locally:
                    missing_locally.remove(filename)  # הקובץ כבר לא חסר בענן

    files_to_delete = [f for f in missing_locally if not f.startswith(".index/")]
    if files_to_delete:
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Warning)
        msg.setWindowTitle("מחיקה מהענן" if not isLTR else "Cloud Delete")
        msg.setText(f"נמצאו {len(files_to_delete)} קבצים למחיקה מהענן. לבצע?" if not isLTR else f"Delete {len(files_to_delete)} files from cloud?")
        msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
        if msg.exec_() == QMessageBox.Ok:
            for filename in files_to_delete:
                delete_from_cloud_with_index(self,filename, prefix=prefix, skip_index_delete = False)
                if filename in missing_locally:
                    missing_locally.remove(filename)  # הקובץ כבר לא חסר בענן

    sync_status = not missing_locally and not missing_in_cloud and not mismatched_files
    sync0 = sync_status
    self.sync0 = sync0

    update_gcs_radio(self)
    if self.sync0:
        self.display_root.setStyleSheet("color: white; background-color: lightblue;")
    else:
        self.display_root.setStyleSheet("color: red; background-color: lightblue;")



    return {
        "missing_in_gcs": missing_in_cloud,
        "missing_locally": missing_locally,
        "mismatched": mismatched_files,
        "sync!": sync_status
    }



def list_s3_contents(bucket, prefix):
    s3 = boto3.client('s3')

    # Ensure prefix ends with / to look INSIDE the folder
    if prefix and not prefix.endswith('/'):
        prefix += '/'
    s3_client= get_s3_client()
    #response = s3.list_objects_v2(
    #    Bucket=bucket,
    #    Prefix=prefix,
    #    Delimiter='/'  # This is the magic "folder" creator
    #)
    response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix, Delimiter='/')
    # 1. Look at Sub-folders (CommonPrefixes)
    print("--- Folders ---")
    for content in response.get('CommonPrefixes', []):
        print(f"Folder: {content.get('Prefix')}")

    # 2. Look at Files (Contents)
    print("\n--- Files ---")
    for obj in response.get('Contents', []):
        if obj['Key'] != prefix:  # Skip the folder object itself
            print(f"File: {obj['Key']} | Size: {obj['Size']} bytes")


def get_aws_cloud_files_recursive(bucket_name, prefix):
    s3_client = get_s3_client()
    cloud_files = {}

    # הכנת הפרמטרים לסריקה
    list_args = {'Bucket': bucket_name}

    # אם יש פריפיקס, נוסיף אותו (מוודא שהוא נגמר ב-/ כדי לסרוק תוכן תיקייה)
    if prefix and prefix.strip():
        search_prefix = prefix if prefix.endswith('/') else f"{prefix}/"
        list_args['Prefix'] = search_prefix

    paginator = s3_client.get_paginator('list_objects_v2')

    try:
        # paginate יסרוק הכל אם אין Prefix, או רק תת-נתיב אם יש
        for page in paginator.paginate(**list_args):
            for obj in page.get('Contents', []):
                full_key = obj['Key']

                # דילוג על שם התיקייה עצמה
                if prefix and full_key == list_args.get('Prefix'):
                    continue

                # סינון קבצים זמניים של אופיס
                if any(x in full_key for x in ['$~', '~$']) or full_key.split('/')[-1].startswith('$'):
                    continue

                try:
                    head = s3_client.head_object(Bucket=bucket_name, Key=full_key)
                    # 1. ניסיון ראשון: המטא-דאטה המותאם אישית שלך (Hex)
                    md5_hash = head.get('Metadata', {}).get('md5-hash')

                    if md5_hash:
                        # אם בטעות שמרת ב-S3 כ-Base64, זה ימיר ל-Hex
                        if '==' in str(md5_hash):
                            import base64
                            md5_hash = base64.b64decode(md5_hash).hex()
                        else:
                            md5_hash = str(md5_hash).strip().lower()

                    cloud_files[full_key] = md5_hash

                except Exception as e:
                    print(f"Error fetching metadata for {full_key}: {e}")

    except Exception as e:
        print(f"Error scanning bucket {bucket_name} (Prefix: {prefix}): {e}")

    return cloud_files

# ==============================================================================
# GCS CORE BROWSER FUNCTION
# ==============================================================================


def browse_gcs_path(self) -> Dict[str, List[str]]:
    """
    Lists the virtual directories (prefixes) and  (blobs) at a given GCS path.
    """
    path_prefix = self.current_path
    # Example usage
    path_prefix = f"{path_prefix}"
    # 1. Normalize the prefix path
    normalized_prefix = path_prefix.strip('/').replace('\\', '/')
    if normalized_prefix and not normalized_prefix.endswith('/'):
        normalized_prefix += '/'


    timer0 = time.time()

    global gcs_client

    # Check 1: If the client already exists, return it instantly (0 seconds)
    if gcs_client is None:
        gcs_client = get_storage_client()

    if gcs_client is None:
        return {"folders": []}



    try:

        bucket = gcs_client.bucket(self.bucket_name)
        blobs_iterator = bucket.list_blobs(prefix=normalized_prefix)

        folders = set()
        timer1 = time.time()
        for blob in blobs_iterator:
            file_path = blob.name

            # 2. Get the path segment AFTER the current normalized_prefix
            relative_path = file_path[len(normalized_prefix):]

            if not relative_path:
                continue

                # 3. Manually find the next level folder
            if '/' in relative_path:
                # If the object path contains a slash, the part before the first slash is the folder name.
                top_level_folder = relative_path.split('/')[0]
                folders.add(top_level_folder)

            # Note: We ignore files at this level (objects without a slash) because the dialog only browses folders.

        final_folders = list(folders)
        time1= timer1-timer0
        time2 = time.time() - timer1
        print(f"DEBUG: Inferred {len(final_folders)} virtual folders successfully. time={time1}:{time2}")

        # Return sorted folders
        return {"folders": final_folders}



    except Exception as e:
        print(f"Error browsing GCS path '{path_prefix}': {e}")
        return {"folders": []}


# ==============================================================================
# GCS BROWSER DIALOG (GUI IMPLEMENTATION)
# ==============================================================================

class GCSBrowserDialog(QtWidgets.QDialog):
    """
    A custom dialog window for navigating GCS virtual directories.
    """

    def __init__(self, parent=None, initial_path=""):
        super().__init__(parent)

        self.provider = parent.provider_info["cloud_provider"]
        self.bucket_name = parent.provider_info["BUCKET_NAME"]
        self.setWindowTitle(f"Browse {self.provider} Bucket")
        self.setGeometry(100, 100, 600, 400)
        self.current_path = initial_path.strip('/').replace('\\', '/')
        self.selected_path = None

        # Widgets
        self.path_label = QtWidgets.QLabel("Path: /")
        self.list_widget = QtWidgets.QListWidget()

        self.up_button = QtWidgets.QPushButton("↑ Up ↑")
        self.up_button.setFixedWidth(120)
        self.up_button.setStyleSheet("padding: 5px 10px; font-size: 18pt;width: 60px; height: 30px;")

        if isLTR:
            self.ok_button = QtWidgets.QPushButton("Select Folder 📂")
            self.cancel_button = QtWidgets.QPushButton("Cancel")
        else:
            self.ok_button = QtWidgets.QPushButton("📂 בחר ספרייה")
            self.cancel_button = QtWidgets.QPushButton("בטל x")


        # Layout
        path_layout = QtWidgets.QHBoxLayout()
        path_layout.addWidget(self.path_label)
        path_layout.addWidget(self.up_button)


        button_layout = QtWidgets.QHBoxLayout()
        button_layout.addStretch(1)
        button_layout.addWidget(self.ok_button)
        button_layout.addWidget(self.cancel_button)

        main_layout = QtWidgets.QVBoxLayout(self)
        main_layout.addLayout(path_layout)
        main_layout.addWidget(self.list_widget)
        main_layout.addLayout(button_layout)

        # Signals
        self.ok_button.clicked.connect(self.accept_selection)
        self.cancel_button.clicked.connect(self.reject)
        self.up_button.clicked.connect(self.go_up_directory)
        self.list_widget.itemDoubleClicked.connect(self.handle_double_click)

        # Initial load
        self.load_directory(self.current_path)


    def load_directory(self, path: str):
        """Fetches contents from GCS and updates the list widget."""
        self.list_widget.clear()

        provider = self.provider
        # Strip trailing slash for display, but keep it internal for logic
        self.current_path = path.strip('/').replace('\\', '/')
        display_path = self.current_path if self.current_path else " (Root)"
        self.path_label.setText(f" {display_path} 📂")
        self.path_label.setStyleSheet("font-size: 20pt; color: black;")

        #contents = browse_gcs_path(self)
        contents = browse_cloud_path(self)

        # 1. Add Folders (with folder icon)
        for folder in contents["folders"]:
            item = QtWidgets.QListWidgetItem(f"{folder}")
            item.setData(QtCore.Qt.UserRole, folder)  # Store the name
            item.setIcon(self.style().standardIcon(QtWidgets.QStyle.SP_DirIcon))
            self.list_widget.addItem(item)



        self.list_widget.sortItems()

    def handle_double_click(self, item: QtWidgets.QListWidgetItem):
        """Handles double-click event: navigate down into a folder."""
        folder_name = item.data(QtCore.Qt.UserRole)

        if folder_name is not None:  # It's a folder, not a file
            new_path = f"{self.current_path}/{folder_name}" if self.current_path else folder_name
            self.load_directory(new_path)

    def go_up_directory(self):
        """Handles the 'Up' button click: navigate one level up."""
        if not self.current_path:
            return  # Already at root

        # Split path and remove the last element
        path_parts = self.current_path.split('/')
        parent_path = '/'.join(path_parts[:-1])

        self.load_directory(parent_path)

    def accept_selection(self):
        """
        Sets the selected path. Prioritizes a highlighted item in the list
        over the current viewing directory (self.current_path).
        """
        selected_items = self.list_widget.selectedItems()

        if selected_items:
            # Case 1: An item (subfolder) is highlighted in the list.
            item = selected_items[0]
            folder_name = item.data(QtCore.Qt.UserRole)

            if folder_name is not None:
                # Construct the path to the SELECTED subfolder
                new_path = f"{self.current_path}/{folder_name}" if self.current_path else folder_name
                self.selected_path = new_path
            else:
                # If the selected item is not a folder (e.g., the 'No Subfolders' text), default to current path.
                self.selected_path = self.current_path
        else:
            # Case 2: Nothing is highlighted, select the current directory (self.current_path).
            self.selected_path = self.current_path

        self.accept()

    @staticmethod
    def get_directory(parent=None, initial_path=""):
        """Static method to show the dialog and return the selected path."""
        dialog = GCSBrowserDialog(parent, initial_path)
        if dialog.exec_() == QtWidgets.QDialog.Accepted:
            return dialog.selected_path
        return None


# ==============================================================================
# STANDALONE TEST ENTRY POINT
# ==============================================================================

if __name__ == "__main__":
    # To run this, ensure environment variables are set:
    # os.environ["amazon_key"] = "..."
    # os.environ["amazon_secret"] = "..."
    # os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "path_to_json"
    BUCKET_NAME=PROVIDER_CONFIG.get("BUCKET_NAME")
    cloud_files = get_gcs_files_recursive(BUCKET_NAME)

    app = QtWidgets.QApplication(sys.argv)

    # Optional: Set a dark theme or style
    app.setStyle("Fusion")
    BUCKET_NAME =""
    print(f" on Bucket: {BUCKET_NAME}") #Starting test with use_mode_aws={use_mode_aws}

    dialog = GCSBrowserDialog()
    if dialog.exec_() == QtWidgets.QDialog.Accepted:
        final_path = dialog.selected_path()
        print(f"USER SELECTED PATH: {final_path}")
        QtWidgets.QMessageBox.information(None, "Selection", f"You chose: {final_path}")
    else:
        print("User cancelled selection.")

    sys.exit(0)