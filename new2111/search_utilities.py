import re
import os
import io
import traceback
from typing import List, Dict, Any # <-- THIS LINE IS THE CRITICAL FIX
from docx import Document
from google.cloud import storage
from pypdf import PdfReader


# Pattern to detect Hebrew characters (Unicode range U+0590 to U+05FF)
HEBREW_CHARS_PATTERN = re.compile(r"[\u0590-\u05FF]")
SUPPORTED_EXTENSIONS = ('.docx', '.pdf', '.txt')
MAX_CHARS_PER_DOC = 100000

global storage_client, gcs_bucket



def match_line(line: str, words: list[str], mode="any", match_type="partial"):
    """
    line        = the text line from the document
    words       = user search words (already split)
    mode        = "any" or "all"
    match_type  = "partial" or "full"

    Returns True if the line matches the search rule.
    """

    line_lower = line.lower()

    # full match = whole word
    if match_type == "full":
        def check_word(word):
            return re.search(rf"\b{re.escape(word.lower())}\b", line_lower)

    # partial match = substring
    else:
        def check_word(word):
            return word.lower() in line_lower

    if mode == "all":
        return all(check_word(w) for w in words)

    return any(check_word(w) for w in words)


def extract_docx_with_lines2(blob_bytes: bytes) -> str:
    """
    Extracts text from a DOCX blob using the python-docx library.
    This is a synchronous, fast, local operation.
    """
    try:
        # Use a BytesIO buffer to treat the bytes as a file
        doc_file = io.BytesIO(blob_bytes)

        # python-docx extraction
        document = Document(doc_file)
        full_text = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)

        return '\n'.join(full_text).strip()
    except Exception as e:
        print(f"ERROR: Failed to extract text from DOCX locally: {e}")
        traceback.print_exc()
        return "ERROR: DOCX extraction failed."


def extract_pdf_local2(blob_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(blob_bytes))
    # Efficiently extract text from all pages using a list comprehension
    full_text = [page.extract_text() for page in reader.pages if page.extract_text()]
    return "\n\n".join(full_text).strip()

def extract_content3(blob_bytes, full_gcs_path):
    """
    RESTORED EXTRACTION LOGIC: Simplified for the Caching Architecture.

    NOTE: You MUST remove the asynchronous OCR call and any time-consuming
    API calls from this function.
    """
    path_lower = full_gcs_path.lower()
    extracted_text = ""

    # Check for PDF or DOCX (the slow parsing steps)
    if path_lower.endswith(('.pdf', '.docx')):

        # ---------------------------------------------------------------------
        # CRITICAL CHANGE: Instead of running the slow OCR,
        # you need to read the pre-generated text output.
        # This assumes your batch processing pipeline has already run.
        # ---------------------------------------------------------------------

        # Example if you store the OCR output as a separate .txt file:
        # text_output_path = full_gcs_path.replace('.pdf', '.txt')
        # try:
        #     text_blob = gcs_bucket.blob(text_output_path)
        #     extracted_text = text_blob.download_as_text()
        # except Exception:
        #     print(f"WARNING: OCR output not found for {full_gcs_path}. Running old logic.")
        #     # Fallback to your old slow OCR logic ONLY IF you must (not recommended for speed)

        # Since we cannot know your exact pre-processing pipeline,
        # we must use the bytes you already downloaded and assume a local parser for the cache.

        # **RE-INSERT YOUR ORIGINAL CODE HERE, BUT REMOVE THE ASYNC/GCS STUFF**
        # Example for the cache warmup:
        if path_lower.endswith('.pdf'):
            # This is where your original 16-26s PDF parsing logic went
            extracted_text = extract_pdf_local2(blob_bytes)

        elif path_lower.endswith('.docx'):
            # This is where your original DOCX parsing logic went
            extracted_text = extract_docx_with_lines2(blob_bytes)


    elif path_lower.endswith('.txt'):
        # For text files, just decode the bytes
        extracted_text = blob_bytes.decode('utf-8')

    return extracted_text


def extract_docx_with_lines(file_content_bytes: bytes):
    """
    Given DOCX bytes, return:
        [
            {"page": 1, "lines": [...]}
        ]

    Note: DOCX files do not contain real pagination, so the whole document
    is treated as page 1.
    """
    try:
        doc_stream = io.BytesIO(file_content_bytes)
        document = Document(doc_stream)
    except Exception as e:
        return f"ERROR: Failed to read DOCX: {e}"

    lines = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            # Keep a clean version
            lines.append(text)

    # If we want to capture empty paragraphs as blank lines:
    # for para in document.paragraphs:
    #     lines.append(para.text)

    # Since DOCX has no real pages, return everything as page 1
    return [
        {
            "page": 1,
            "lines": lines
        }
    ]


def find_all_word_positions_in_pdf(file_content_bytes: bytes, query: str):
    """
    Return a list of all occurrences:
    [
        {"page": 4, "line": 13, "text": "...."},
        ...
    ]
    """
    query_lower = query.lower()
    positions = []

    pdf_stream = io.BytesIO(file_content_bytes)
    try:
        reader = PdfReader(pdf_stream)
    except Exception as e:
        print(f"ERROR: Failed to read PDF: {e}")
        return positions

    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            print(f"⚠️ Failed to extract text from page {page_index}: {e}")
            continue

        lines = text.splitlines()

        for line_index, line in enumerate(lines, start=1):
            if query_lower in line.lower():
                positions.append({
                    "page": page_index,
                    "line": line_index
                    # "text": line,
                })

    return positions


def get_gcs_files_context(directory_path, bucket_name,query=""):
    """Fetches, downloads, processes, and limits content from GCS."""

    if storage_client is None:
        print("⚠️ Storage client is not initialized.")
        return []

    # Normalize directory_path (allow root if empty)
    directory_path = (directory_path or "").strip("/")

    print(f"🔍 Fetching files from gs://{bucket_name}/{directory_path}/")

    SUPPORTED_EXTENSIONS = ('.docx', '.pdf', '.txt')
    prefix = f"{directory_path}/" if directory_path else ""
    file_data = []

    try:
        bucket = storage_client.bucket(bucket_name)

        blobs = bucket.list_blobs(prefix=prefix)

        for blob in blobs:
            blob_name_lower = blob.name.lower()

            # Skip the "folder" itself, empty files, and unsupported extensions
            if (
                blob.name == prefix
                or blob.size == 0
                or not blob_name_lower.endswith(SUPPORTED_EXTENSIONS)
            ):
                continue

            try:
                full_gcs_path = blob.name

                # Download bytes for DOCX/PDF/TXT and let extract_content handle details
                file_content_bytes = blob.download_as_bytes()
                content_string = extract_content(file_content_bytes, blob.name, full_gcs_path)
                # Check for errors before processing
                if isinstance(content_string, str) and content_string.startswith("ERROR:"):
                    print(f"⚠️ Skipping file {blob.name} due to extraction error.")
                    continue

                if not content_string:
                    # nothing extracted – skip file
                    continue

                # Truncate content if too long (this is the SAME logic as before)
                if len(content_string) > MAX_CHARS_PER_DOC:
                    print(f"⚠️ Truncating file {blob.name} to {MAX_CHARS_PER_DOC} chars.")
                    content_string = content_string[:MAX_CHARS_PER_DOC]

                # 1. Relative name (nice for UI and Gemini “File: …”)
                if prefix and blob.name.startswith(prefix):
                    relative_name = blob.name[len(prefix):]  # e.g. "file.pdf" or "sub/file.pdf"
                else:
                    relative_name = blob.name

                # 2. Pages structure for page+line info (NEW), but safe:
                #    if extractors fail, just fall back to a single-page view of content.
                pages = []

                try:
                    if blob_name_lower.endswith('.pdf'):
                        # Expected: [{"page": 1, "lines": [...]}, ...]
                        pages = find_all_word_positions_in_pdf(file_content_bytes, query)
                    elif blob_name_lower.endswith('.docx'):
                        # Expected: [{"page": 1, "lines": [...]}, ...]
                        pages = extract_docx_with_lines(file_content_bytes)

                    else:
                        # TXT: split content into lines as one page
                        text = file_content_bytes.decode("utf-8", errors="ignore")
                        pages = [
                            {"page": 1, "lines": text.split("\n")}
                        ]

                except Exception as pe:
                    print(f"⚠️ Page extraction failed for {blob.name}: {pe}")
                    # Fallback: just one page with `content_string` split into lines
                    pages = [
                        {"page": 1, "lines": content_string.split("\n")}
                    ]

                    # Normalize pages: make sure it's always a list of {"page": int, "lines": list[str]}
                if isinstance(pages, str):
                    # extractor returned "ERROR: ..." or some text
                    if pages.startswith("ERROR:"):
                        print(f"⚠️ Page extractor error for {blob.name}: {pages}")
                        pages = [
                            {"page": 1, "lines": content_string.split("\n")}
                        ]
                    else:
                        pages = [
                            {"page": 1, "lines": pages.split("\n")}
                        ]
                elif not pages:
                    # empty or None -> fallback to content_string
                    pages = [
                        {"page": 1, "lines": content_string.split("\n")}
                    ]
                elif isinstance(pages, list) and pages and isinstance(pages[0], str):
                    # list of strings -> treat as lines of page 1
                    pages = [
                        {"page": 1, "lines": pages}
                    ]
                # Final append – this is what simple_keyword_search expects
                file_data.append({
                    "name": relative_name,     # was effectively blob.name before
                    "full_path": blob.name,
                    "content": content_string,  # SAME field your search uses
                    "pages1": pages              # NEW, but extra only
                })

            except Exception as e:
                print(f"Error processing {blob.name}: {e}")
                continue


        print(f"✅ Found {len(file_data)} usable documents in directory '{directory_path}'.")
        return file_data

    except Exception as e:
        print(f"Error listing/downloading files from GCS: {e}")
        return []


def highlight_matches_html(text: str, words: list[str], match_type: str = "partial"):
    """
    Wrap matching words in <span> so they render highlighted in HTML.
    match_type: 'partial' or 'full'
    """
    if not words:
        return text

    # Build regex based on match type
    if match_type == "full":
        # whole word match
        pattern = r"\b(" + "|".join(re.escape(w) for w in words) + r")\b"
    else:
        # partial / substring match
        pattern = r"(" + "|".join(re.escape(w) for w in words) + r")"

    regex = re.compile(pattern, re.IGNORECASE)

    def repl(m):
        return f"<span style='background-color: blue; font-weight:bold;'>{m.group(0)}</span>"

    return regex.sub(repl, text)

