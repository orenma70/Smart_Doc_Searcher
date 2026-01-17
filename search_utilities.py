from google import genai
import os
import io
import time
from google.cloud import storage
from google.cloud import vision_v1 as vision
import traceback
from pathlib import Path
import threading
from concurrent.futures import ThreadPoolExecutor
from typing import List, Dict, Any, Optional
import json
from docx import Document
from document_parsers import extract_text_and_images_from_pdf
import pytesseract
from PIL import Image
from azure.storage.blob import BlobServiceClient




# --- Global Client Variables (Set to None for Lazy Loading) ---
storage_client: Optional[storage.Client] = None
vision_client: Optional[storage.Client] = None
gemini_client: Optional[storage.Client] = None
gcs_bucket: Optional[storage.Client] = None  # <-- NEW: Store the GCS Bucket object once

# --- Global Shared State for Caching ---

IS_CLOUD_RUN = bool(os.environ.get('K_SERVICE'))
MAX_CHARS_PER_DOC = 100000
MAX_WAIT_SECONDS = 600
POLL_INTERVAL_SECONDS = 10
MAX_CONCURRENT_DOWNLOADS = 10


# In search_utilities.py (UPDATED GLOBALS)
# --- Global Shared State for Caching ---
# Key: directory_path (e.g., "my_folder/docs")
# Value: List[Dict] (The document data for that path)
DIRECTORY_CACHE_MAP: Dict[str, List[Dict]] = {}
cache_lock = threading.Lock() # Use this lock for thread-safe access to the map



def get_hd_files_context(directory_path: str, local_root: str) -> List[Dict[str, Any]]:
    """
    Simulates GCS file fetching by recursively reading files from the local hard drive.

    Args:
        directory_path: The path requested by the user (e.g., 'archive/2025').
        local_root: The physical root directory on the local disk.

    Returns:
        A list of document dictionaries in the same format as the GCS fetch.
    """

    # 1. Calculate the full physical path to start the search
    # directory_path (e.g., 'archive/2025') is appended to the local root
    search_root = Path(local_root) / directory_path.strip("/")

    if not search_root.is_dir():
        print(f"ERROR: Local search directory not found: {search_root}")
        return []

    documents_list = []

    # 2. Use os.walk for recursive file listing (simulating GCS recursion)
    # os.walk yields (root, dirs, files) for each directory
    for root, _, files in os.walk(search_root):
        root_path = Path(root)

        for file_name in files:
            full_file_path = (root_path / file_name)

            # 3.
            # This ensures the 'full_path' key matches the GCS/Cache format (e.g., 'archive/2025/doc.txt')
            try:
                # Use .as_posix() to ensure forward slashes, even on Windows
                relative_path_for_cache = full_file_path.relative_to(local_root).as_posix()
            except ValueError:
                # Should not happen if paths are correct, but handles files outside the root
                continue

                # 4. Read File Content (Equivalent to GCS blob download)
            # You would use your real file parsing logic here (PDF, DOCX, etc.)
            # For simplicity, we assume text files for debugging.
            try:
                # IMPORTANT: Use 'rb' (read binary) for compatibility if you later parse binary files
                # like PDFs or DOCX, then process content inside your actual parser.
                # For basic text reading:
                file_ext = full_file_path.suffix.lower()
                content = ""

                if file_ext == '.pdf':
                    content = extract_text_and_images_from_pdf(str(full_file_path))
                    content = content[0]
                elif file_ext in ['.txt', '.csv']:
                    # Read text files normally
                    with open(full_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                elif file_ext in ['.docx', '.pptx']:
                    content = ""
                    doc1 = Document(str(full_file_path))
                    for paragraph in doc1.paragraphs:
                        content += paragraph.text + "\n"
                    content=content.strip()


                else:
                    content = ""  # Skip unsupported files

                documents_list.append({
                    "name": file_name,
                    "full_path": relative_path_for_cache.lower(),  # Store as lowercase for cache filter safety
                    "content": content,
                    "source": f"Local HD: {relative_path_for_cache}",
                    "pages": [{"page": 1, "lines": content.split("\n")}]
                    # You may include more metadata (like size, last modified) here
                })
            except Exception as e:
                # Log issues reading specific files
                print(f"Warning: Could not process file {full_file_path}: {e}")

    return documents_list


# In search_utilities.py
# In search_utilities.py

def put_documents_in_cache(path_key: str, documents: List[Dict[str, Any]]):
    """Safely adds a list of documents to the global cache under the given key."""
    global DIRECTORY_CACHE_MAP, cache_lock

    # Ensure the key is fully normalized (lowercase, no slashes)
    normalized_key = path_key.strip("/").lower()

    with cache_lock:
        if documents:
            DIRECTORY_CACHE_MAP[normalized_key] = documents
            print(f"CACHE-PUT: Stored {len(documents)} documents for '{normalized_key}'.")


def get_documents_from_cache(directory_path: str) -> Optional[List[Dict[str, Any]]]:
    """
    Attempts a hierarchy-aware lookup for documents, checking the exact path
    and then searching parent paths. Returns None on cache miss.
    """
    global DIRECTORY_CACHE_MAP, cache_lock

    cleaned_path = directory_path.strip("/").lower()

    with cache_lock:

        # 1a. EXACT MATCH
        if cleaned_path in DIRECTORY_CACHE_MAP:
            print(f"CACHE-GET: Exact hit for '{cleaned_path}'.")
            return list(DIRECTORY_CACHE_MAP[cleaned_path])

        # 1b. PARENT/HIERARCHY MATCH
        path_segments = cleaned_path.split('/')
        potential_parent_keys = []

        # Generate parent keys (e.g., 'main_dir/dir1' -> 'main_dir')
        for i in range(len(path_segments) - 1, 0, -1):
            key = '/'.join(path_segments[:i])
            if key:
                potential_parent_keys.append(key)
        potential_parent_keys.append("")  # Check the root key

        for parent_key in potential_parent_keys:
            if parent_key in DIRECTORY_CACHE_MAP:
                print(f"CACHE-GET: Parent hit on '{parent_key}' for '{cleaned_path}'.")

                # --- Filtering ---
                prefix_with_slash = cleaned_path + '/'

                filtered_documents = [
                    doc for doc in DIRECTORY_CACHE_MAP[parent_key]
                    if doc.get('full_path', '').lower().startswith(prefix_with_slash)
                ]

                if filtered_documents:
                    return list(filtered_documents)

    # Cache Miss
    return None


def get_documents_for_path(self,directory_path: str) -> List[Dict[str, Any]]:
    # 1. נרמול הנתיב (חשוב להתאמה למפתחות בענן)
    cleaned_path = directory_path.strip("/")

    # 2. ניסיון משיכה מה-Cache (נשאר ללא שינוי - FAST PATH)
    cached_docs = get_documents_from_cache(cleaned_path)
    if cached_docs is not None:
        return cached_docs

    # 3. CACHE MISS - שליפה מה-Bucket (SLOW PATH)
    print(f"CACHE-MISS: Fetching JSON indices for '{cleaned_path}'.")

    fetched_documents = []

    try:
        client = get_storage_client()
        bucket = client.bucket(self.provider_info["BUCKET_NAME"])

        # הגדרת ה-prefix לחיפוש הקבצים המקוריים
        prefix = cleaned_path + "/" if cleaned_path else ""
        blobs = bucket.list_blobs(prefix=prefix)

        for blob in blobs:
            key = blob.name
            # דילוג על תיקיות והתיקייה .index עצמה
            if key.endswith('/') or key.startswith('.index/'):
                continue

            # בניית הנתיב לקובץ ה-JSON התואם
            base_name = os.path.splitext(key)[0]
            index_key = f".index/{base_name}.json"

            try:
                index_blob = bucket.blob(index_key)
                if index_blob.exists():
                    # הורדת ה-JSON וטעינת הנתונים
                    json_data = json.loads(index_blob.download_as_text())

                    # בניית האובייקט שהחיפוש מצפה לו
                    fetched_documents.append({
                        "name": os.path.basename(key),
                        "full_path": key,
                        "pages": json_data.get("pages", []),
                        # יצירת שדה content שטוח לטובת ה-Gemini (RAG)
                        "content": "\n".join([" ".join(p.get("lines", [])) for p in json_data.get("pages", [])])
                    })
            except Exception as e:
                print(f"Skipping {key}: Index error: {e}")

    except Exception as e:
        print(f"🔥 GCS Fetch Error: {e}")

    # 4. שמירה ב-Cache (כדי שהחיפוש הבא יהיה מהיר)
    if fetched_documents:
        put_documents_in_cache(cleaned_path, fetched_documents)

    return fetched_documents

def initialize_all_clients(self):
    """
    Initializes clients only if they haven't been initialized yet.
    This runs inside the request context on the first call.
    """
    # CRITICAL: Add gcs_bucket to the global list
    global storage_client, gemini_client, vision_client, gcs_bucket

    # 1. Initialize Storage Client and Bucket
    if storage_client is None:
        storage_client = get_storage_client()

        # --- CRITICAL INSERTION POINT ---
        if storage_client is not None and self.provider_info["BUCKET_NAME"] and gcs_bucket is None:
            try:
                # This line sets the global gcs_bucket variable!
                gcs_bucket = storage_client.bucket(self.provider_info["BUCKET_NAME"])
                print(f"✅ GCS STEP 2: Successfully connected to Bucket '{self.provider_info["BUCKET_NAME"]}'.")
            except Exception as e:
                print(f"FATAL: GCS STEP 2: FAILED to get Bucket '{self.provider_info["BUCKET_NAME"]}'. Check IAM Permissions. Error: {e}")
                gcs_bucket = None  # Ensure it is None on failure
        # --- END CRITICAL INSERTION POINT ---

    # 2. Initialize other clients
    if vision_client is None:
        vision_client = get_vision_client()

    if gemini_client is None:
        gemini_client = get_gemini_client()

    # Returns True only if all critical clients are initialized
    return (storage_client is not None and gcs_bucket is not None and  # Ensure gcs_bucket is checked!
            gemini_client is not None and vision_client is not None)


def get_storage_client():
    """Initializes and returns the Google Cloud Storage client."""
    try:
        return storage.Client()
    except Exception as e:
        print(f"FATAL: Could not initialize GCS client. Error: {e}")
        return None


def get_vision_client():
    """Initializes and returns the Google Cloud Vision client."""
    try:
        # return vision.V1.ImageAnnotatorClient() #vision.ImageAnnotatorClient()
        return vision.ImageAnnotatorClient()
    except Exception as e:
        print(f"FATAL: Could not initialize Vision client. Error: {e}")
        return None


def get_gemini_client():
    """Initializes and returns the Gemini client."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        print("FATAL: GEMINI_API_KEY environment variable not set.")
        return None
    try:
        return genai.Client(api_key=api_key)
    except Exception as e:
        print(f"FATAL: Could not initialize Gemini client. Check API Key. Error: {e}")
        return None


# Add this function to search_utilities.py
def get_gcs_bucket(self):
    """Returns the initialized GCS bucket object, ensuring initialization is attempted."""
    global gcs_bucket
    if gcs_bucket is None:
        initialize_all_clients(self)
    return gcs_bucket


# Add this function (or just use get_gcs_bucket logic)
def get_storage_client_instance(self):
    """Returns the initialized Storage client object."""
    global storage_client
    if storage_client is None:
        initialize_all_clients(self)
    return storage_client


# Add this function
def get_vision_client_instance(self):
    """Returns the initialized Vision client object."""
    global vision_client
    if vision_client is None:
        initialize_all_clients(self)
    return vision_client


# Add this function
def get_gemini_client_instance(self):
    """Returns the initialized Gemini client object."""
    global gemini_client
    if gemini_client is None:
        initialize_all_clients(self)
    return gemini_client



def detect_text_gcs_async(self,gcs_uri, gcs_destination_uri):
    """
    Performs asynchronous OCR on a PDF in GCS using the Vision API.
    gcs_uri: Input GCS path of the PDF (e.g. 'gs://bucket/folder/file.pdf').
    gcs_destination_uri: GCS folder prefix where JSON results will be written
                         (e.g. 'gs://bucket/vision_ocr_output/job123/').
    """
    vision_client = get_vision_client_instance(self)

    if vision_client is None:
        return "ERROR: Vision client not initialized."

    print(f"👁️ Starting ASYNC PDF OCR for {gcs_uri} -> {gcs_destination_uri}")

    # 1) Build Vision request objects
    feature = vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)

    gcs_source = vision.GcsSource(uri=gcs_uri)
    input_config = vision.InputConfig(
        gcs_source=gcs_source,
        mime_type="application/pdf",  # ✅ important
    )

    gcs_destination = vision.GcsDestination(uri=gcs_destination_uri)
    output_config = vision.OutputConfig(
        gcs_destination=gcs_destination,
        batch_size=5,  # pages per JSON file (tune if needed)
    )

    async_request = vision.AsyncAnnotateFileRequest(
        features=[feature],
        input_config=input_config,
        output_config=output_config,
    )

    # 2) Call async_batch_annotate_files (this is the correct method)
    operation = vision_client.async_batch_annotate_files(requests=[async_request])

    try:
        result = operation.result(timeout=300)  # wait up to 5 minutes
        print("✅ ASYNC PDF OCR operation finished.")
    except Exception as e:
        print(f"❌ ASYNC OCR operation failed or timed out: {e}")
        traceback.print_exc()
        return f"ERROR: ASYNC OCR failed or timed out: {e}"

    # 3) Read JSON outputs from the destination bucket
    full_text = ""
    storage_cli = get_storage_client_instance()

    # Parse bucket and prefix from 'gs://bucket/prefix/...'
    parts = gcs_destination_uri.replace("gs://", "").split("/", 1)
    out_bucket_name = parts[0]
    out_prefix = parts[1].rstrip("/") + "/"
    print(f"✅ debug2 list_blobs .{out_prefix}")
    try:
        bucket = storage_cli.bucket(out_bucket_name)
        blobs = bucket.list_blobs(prefix=out_prefix)

        for blob in blobs:
            if not blob.name.endswith(".json"):
                continue

            json_bytes = blob.download_as_bytes()
            json_data = json.loads(json_bytes)

            # Each JSON file corresponds to an AnnotateFileResponse with 'responses'
            for response in json_data.get("responses", []):
                full_annotation = response.get("fullTextAnnotation")
                if full_annotation and full_annotation.get("text"):
                    full_text += full_annotation["text"] + "\n"

            # Optional cleanup: delete JSON result files to avoid clutter
            blob.delete()

        if full_text.strip():
            print("✅ ASYNC PDF OCR successful, text extracted.")
            return full_text
        else:
            print("⚠️ ASYNC PDF OCR executed but found no text.")
            return "ERROR: ASYNC PDF OCR executed but found no text."

    except Exception as e:
        print(f"❌ Error reading OCR results or cleaning up: {e}")
        traceback.print_exc()
        return f"ERROR: Error processing OCR output: {e}"



def extract_content(self,blob_bytes, blob_name, full_gcs_path):
    """
    Extracts text content from document bytes.

    - PDF files are routed to the asynchronous Vision API OCR (detect_text_gcs_async).
    - DOCX files use the docx library.
    - Other files are treated as plain text.
    """

    blob_name_lower = blob_name.lower()
    prefix = "gs://" + self.provider_info["BUCKET_NAME"] + "/"
    gcs_uri = prefix + full_gcs_path

    # Check for PDF
    if blob_name_lower.endswith('.pdf'):
        # NEW LOGIC: Use ASYNCHRONOUS Vision OCR for all PDFs
        # (This handles both scanned images and selectable text robustly)

        # Create a unique temporary path for this job's output
        job_id = f"ocr_job_{os.path.basename(full_gcs_path)}_{int(time.time())}"
        gcs_destination_uri = self.provider_info["GCS_OCR_OUTPUT_PATH"] + job_id + "/"

        # Call the async OCR function defined in step 2
        try:
            content = detect_text_gcs_async(gcs_uri, gcs_destination_uri)
            return content
        except NameError:
            # Fallback if the function is not defined, useful for debugging
            return "ERROR: detect_text_gcs_async function is not defined."

    # Check for DOCX
    elif blob_name_lower.endswith('.docx'):
        text = ""
        try:
            document = Document(io.BytesIO(blob_bytes))
            combined_text = []

            # 1. Extract Regular Editable Text
            for paragraph in document.paragraphs:
                combined_text.append(paragraph.text)

            # 2. Extract and OCR Embedded Images

            for rel in document.part.rels.items():
                target_ref_safe = getattr(rel, 'target_ref', '')
                if target_ref_safe:
                    try:
                        if "image" in target_ref_safe or "/media/" in target_ref_safe:
                            image_part = rel.target_part
                            image_bytes = image_part.blob
                            image = Image.open(io.BytesIO(image_bytes))

                            # Run OCR
                            ocr_text = pytesseract.image_to_string(image, 'heb')
                            # a = b[4]
                            # Append OCR results with a separator/label
                            combined_text.append(ocr_text)
                    except pytesseract.TesseractNotFoundError:
                        # IMPORTANT: This error means Tesseract executable is missing in the cloud!
                        combined_text.append("\n[OCR ERROR: Tesseract not found on system path. Image skipped.]")
                    except Exception as e:
                        # Handle other potential errors (like missing image dependencies)
                        combined_text.append(f"\n[OCR ERROR: Failed to process image: {e}]")

            # 3. Combine and Optionally Search
            text = "\n".join(combined_text)

            #for paragraph in document.paragraphs:
            #    text += paragraph.text + "\n"
            return text.strip()

        except Exception as e:
            return f"ERROR: Could not read DOCX content: {e}"

    # Handle other files as plain text (TXT, CSV, etc.)
    try:
        # We assume the bytes are UTF-8 encoded text
        return blob_bytes.decode('utf-8', errors='ignore').strip()
    except Exception as e:
        return f"ERROR: Could not decode text content for {blob_name}. {e}"



def get_gcs_files_context(self,directory_path: str, bucket_name: str, query: str = "") -> List[Dict[str, Any]]:
    """
    Fetches, downloads, and processes content from GCS concurrently using a ThreadPoolExecutor.
    This prevents the request from hitting a 504 Gateway Timeout on large directories.
    """
    directory_path = (directory_path or "").strip("/")
    prefix = f"{directory_path}/" if directory_path else ""
    file_data: List[Dict[str, Any]] = []
    storage_client = get_storage_client_instance(self)
    # 1. Get the list of blobs (metadata only - this is fast)
    print(f"prefix: debug to list GCS blobs: {prefix}")
    try:
        bucket = storage_client.bucket(bucket_name)
        blobs = bucket.list_blobs(prefix=prefix)
        # Filter for relevant file types immediately
        blobs_to_process = [
            blob for blob in blobs
            if blob.name.endswith(('.docx', '.pdf', '.txt'))
        ]

    except Exception as e:
        print(f"ERROR: Failed to list GCS blobs: {e}")
        return []

    def process_single_blob(blob):
        """Worker function executed by each thread."""
        try:
            # NETWORK I/O: This is the slow part now happening in parallel
            file_content_bytes = blob.download_as_bytes()
            content_string = extract_content(file_content_bytes, blob.name, blob.name)

            if content_string.startswith("ERROR:") or not content_string:
                return None

            # Determine relative name
            relative_name = blob.name[len(prefix):] if prefix and blob.name.startswith(prefix) else blob.name
            normalized_gcs_path = Path(blob.name).as_posix()
            print(f"full path={normalized_gcs_path}")
            return {
                "name": relative_name,
                "full_path": normalized_gcs_path,
                "content": content_string[:MAX_CHARS_PER_DOC],  # Assuming MAX_CHARS_PER_DOC is global
                # Pages structure is often too complex to compute concurrently,
                # but adding a placeholder for consistency:
                "pages": [{"page": 3, "lines": content_string.split("\n")}],
            }
        except Exception as e:
            print(f"ERROR processing {blob.name}: {e}")
            return None

    # 2. Execute file processing concurrently
    with ThreadPoolExecutor(max_workers=MAX_CONCURRENT_DOWNLOADS) as executor:
        # Map the worker function over the list of blobs
        results = executor.map(process_single_blob, blobs_to_process)

        # Collect results, filtering out None (failed/skipped files)
        file_data = [res for res in results if res is not None]

    return file_data

# ==============================================================================
# --- MOCK & PLACEHOLDER FUNCTIONS ---
# NOTE: Replace these with your actual implementations
# ==============================================================================

# ==============================================================================
# --- REVISED SEARCH FUNCTION ---
# ==============================================================================
