import re
import os
import io
import traceback
from typing import List, Dict, Any # <-- THIS LINE IS THE CRITICAL FIX
from docx import Document
from google.cloud import storage
from pypdf import PdfReader


# Pattern to detect Hebrew characters (Unicode range U+0590 to U+05FF)
HEBREW_CHARS_PATTERN = re.compile(r"[\u0590-\u05FF]")
SUPPORTED_EXTENSIONS = ('.docx', '.pdf', '.txt')
MAX_CHARS_PER_DOC = 100000

global storage_client, gcs_bucket

def split_into_sentences(text: str) -> List[str]:
    """
    Splits text into segments based purely on sentence endings (., ?, !)
    to create logical 'lines' or segments, especially useful for giant text blocks.
    """
    sentences = re.split(r'([.?!])\s*(?=[A-Z\u0590-\u05FF]|$)', text.strip())

    segments = []
    current_segment = ""

    for part in sentences:
        if not part.strip():
            continue

        if part in ['.', '?', '!']:
            current_segment += part
            if current_segment.strip():
                segments.append(current_segment.strip())
            current_segment = ""
            continue

        if current_segment:
            current_segment += (" " + part.strip())
        else:
            current_segment = part.strip()

    if current_segment:
        segments.append(current_segment.strip())

    return segments


def split_into_paragraphs(text: str) -> List[str]:
    """
    Splits text into paragraphs, with advanced fallback for poor parser results.
    """

    lines = text.split("\n")

    if len(lines) < 10 and len(text) > 500:
        paragraphs = split_into_sentences(text)
        if len(paragraphs) > 1:
            return paragraphs

    paragraphs = []
    current = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            continue
        current.append(stripped)
    if current:
        paragraphs.append(" ".join(current))

    return paragraphs


def match_line(text: str, words: List[str], mode: str = "any", match_type: str = "partial") -> bool:
    """Checks if a line/paragraph matches the query words (BiDi safe)."""
    text_lower = text.lower()
    patterns_to_check = words + [w[::-1] for w in words]
    patterns_to_check = [p.strip() for p in patterns_to_check if p.strip()]

    if match_type == "full":
        def check_word(word_list):
            return [re.search(rf"\b{re.escape(w.lower())}\b", text_lower) for w in word_list]
    else:
        def check_word(word_list):
            return [w.lower() in text_lower for w in word_list]

    all_matches_found = [m for w in patterns_to_check for m in check_word([w]) if m]

    if mode == "all":
        match_count = 0
        for w in words:
            if w.strip():
                if (match_type == "full" and (check_word([w])[0] or check_word([w[::-1]])[0])) or \
                        (match_type == "partial" and (w.lower() in text_lower or w[::-1].lower() in text_lower)):
                    match_count += 1
        return match_count == len(words)

    return any(all_matches_found)

def extract_docx_with_lines2(blob_bytes: bytes) -> str:
    """
    Extracts text from a DOCX blob using the python-docx library.
    This is a synchronous, fast, local operation.
    """
    try:
        # Use a BytesIO buffer to treat the bytes as a file
        doc_file = io.BytesIO(blob_bytes)

        # python-docx extraction
        document = Document(doc_file)
        full_text = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)

        return '\n'.join(full_text).strip()
    except Exception as e:
        print(f"ERROR: Failed to extract text from DOCX locally: {e}")
        traceback.print_exc()
        return "ERROR: DOCX extraction failed."


def extract_pdf_local2(blob_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(blob_bytes))
    # Efficiently extract text from all pages using a list comprehension
    full_text = [page.extract_text() for page in reader.pages if page.extract_text()]
    return "\n\n".join(full_text).strip()

def extract_content3(blob_bytes, full_gcs_path):
    """
    RESTORED EXTRACTION LOGIC: Simplified for the Caching Architecture.

    NOTE: You MUST remove the asynchronous OCR call and any time-consuming
    API calls from this function.
    """
    path_lower = full_gcs_path.lower()
    extracted_text = ""

    # Check for PDF or DOCX (the slow parsing steps)
    if path_lower.endswith(('.pdf', '.docx')):

        # ---------------------------------------------------------------------
        # CRITICAL CHANGE: Instead of running the slow OCR,
        # you need to read the pre-generated text output.
        # This assumes your batch processing pipeline has already run.
        # ---------------------------------------------------------------------

        # Example if you store the OCR output as a separate .txt file:
        # text_output_path = full_gcs_path.replace('.pdf', '.txt')
        # try:
        #     text_blob = gcs_bucket.blob(text_output_path)
        #     extracted_text = text_blob.download_as_text()
        # except Exception:
        #     print(f"WARNING: OCR output not found for {full_gcs_path}. Running old logic.")
        #     # Fallback to your old slow OCR logic ONLY IF you must (not recommended for speed)

        # Since we cannot know your exact pre-processing pipeline,
        # we must use the bytes you already downloaded and assume a local parser for the cache.

        # **RE-INSERT YOUR ORIGINAL CODE HERE, BUT REMOVE THE ASYNC/GCS STUFF**
        # Example for the cache warmup:
        if path_lower.endswith('.pdf'):
            # This is where your original 16-26s PDF parsing logic went
            extracted_text = extract_pdf_local2(blob_bytes)

        elif path_lower.endswith('.docx'):
            # This is where your original DOCX parsing logic went
            extracted_text = extract_docx_with_lines2(blob_bytes)


    elif path_lower.endswith('.txt'):
        # For text files, just decode the bytes
        extracted_text = blob_bytes.decode('utf-8')

    return extracted_text


def extract_docx_with_lines(file_content_bytes: bytes):
    """
    Given DOCX bytes, return:
        [
            {"page": 1, "lines": [...]}
        ]

    Note: DOCX files do not contain real pagination, so the whole document
    is treated as page 1.
    """
    try:
        doc_stream = io.BytesIO(file_content_bytes)
        document = Document(doc_stream)
    except Exception as e:
        return f"ERROR: Failed to read DOCX: {e}"

    lines = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            # Keep a clean version
            lines.append(text)

    # If we want to capture empty paragraphs as blank lines:
    # for para in document.paragraphs:
    #     lines.append(para.text)

    # Since DOCX has no real pages, return everything as page 1
    return [
        {
            "page": 1,
            "lines": lines
        }
    ]


def find_all_word_positions_in_pdf(file_content_bytes: bytes, query: str):
    """
    Return a list of all occurrences:
    [
        {"page": 4, "line": 13, "text": "...."},
        ...
    ]
    """
    query_lower = query.lower()
    positions = []

    pdf_stream = io.BytesIO(file_content_bytes)
    try:
        reader = PdfReader(pdf_stream)
    except Exception as e:
        print(f"ERROR: Failed to read PDF: {e}")
        return positions

    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            print(f"⚠️ Failed to extract text from page {page_index}: {e}")
            continue

        lines = text.splitlines()

        for line_index, line in enumerate(lines, start=1):
            if query_lower in line.lower():
                positions.append({
                    "page": page_index,
                    "line": line_index
                    #"text": line,
                })

    return positions


# In search_utilities.py, replace the start of this function:
def get_gcs_files_context(directory_path: str, bucket_name: str) -> Dict[str, Dict[str, Any]]:
    """
    Lists, downloads, and extracts content from GCS synchronously.
    CRITICAL: Initializes its own storage.Client() to be self-contained.
    """
    print(f"INFO: Running synchronous GCS fetch for path: {directory_path}")
    try:
        # --- THIS BLOCK MUST BE PRESENT AND LOCAL ---
        storage_client = storage.Client()  # This defines storage_client locally
        bucket = storage_client.bucket(bucket_name)
        # -------------------------------------------
    except Exception as e:
        print(f"ERROR: Could not initialize GCS client or bucket in legacy search: {e}")
        traceback.print_exc()
        return {}



    prefix = f"{directory_path}/" if directory_path else ""
    file_data = []
    try:
        bucket = storage_client.bucket(bucket_name)

        blobs = bucket.list_blobs(prefix=prefix)

        for blob in blobs:
            blob_name_lower = blob.name.lower()

            # Skip the "folder" itself, empty files, and unsupported extensions
            if (
                blob.name == prefix
                or blob.size == 0
                or not blob_name_lower.endswith(SUPPORTED_EXTENSIONS)
            ):
                continue

            try:
                full_gcs_path = blob.name

                # Download bytes for DOCX/PDF/TXT and let extract_content handle details
                file_content_bytes = blob.download_as_bytes()
                content_string = extract_content3(file_content_bytes, blob.name, full_gcs_path)
                # Check for errors before processing
                if isinstance(content_string, str) and content_string.startswith("ERROR:"):
                    print(f"⚠️ Skipping file {blob.name} due to extraction error.")
                    continue

                if not content_string:
                    # nothing extracted – skip file
                    continue

                # Truncate content if too long (this is the SAME logic as before)
                if len(content_string) > MAX_CHARS_PER_DOC:
                    print(f"⚠️ Truncating file {blob.name} to {MAX_CHARS_PER_DOC} chars.")
                    content_string = content_string[:MAX_CHARS_PER_DOC]

                # 1. Relative name (nice for UI and Gemini “File: …”)
                if prefix and blob.name.startswith(prefix):
                    relative_name = blob.name[len(prefix):]  # e.g. "file.pdf" or "sub/file.pdf"
                else:
                    relative_name = blob.name

                # 2. Pages structure for page+line info (NEW), but safe:
                #    if extractors fail, just fall back to a single-page view of content.
                pages = []

                try:
                    if blob_name_lower.endswith('.pdf'):
                        # Expected: [{"page": 1, "lines": [...]}, ...]
                        pages = find_all_word_positions_in_pdf(file_content_bytes, query)
                    elif blob_name_lower.endswith('.docx'):
                        # Expected: [{"page": 1, "lines": [...]}, ...]
                        pages = extract_docx_with_lines(file_content_bytes)

                    else:
                        # TXT: split content into lines as one page
                        text = file_content_bytes.decode("utf-8", errors="ignore")
                        pages = [
                            {"page": 1, "lines": text.split("\n")}
                        ]

                except Exception as pe:
                    print(f"⚠️ Page extraction failed for {blob.name}: {pe}")
                    # Fallback: just one page with `content_string` split into lines
                    pages = [
                        {"page": 1, "lines": content_string.split("\n")}
                    ]

                    # Normalize pages: make sure it's always a list of {"page": int, "lines": list[str]}
                if isinstance(pages, str):
                    # extractor returned "ERROR: ..." or some text
                    if pages.startswith("ERROR:"):
                        print(f"⚠️ Page extractor error for {blob.name}: {pages}")
                        pages = [
                            {"page": 1, "lines": content_string.split("\n")}
                        ]
                    else:
                        pages = [
                            {"page": 1, "lines": pages.split("\n")}
                        ]
                elif not pages:
                    # empty or None -> fallback to content_string
                    pages = [
                        {"page": 1, "lines": content_string.split("\n")}
                    ]
                elif isinstance(pages, list) and pages and isinstance(pages[0], str):
                    # list of strings -> treat as lines of page 1
                    pages = [
                        {"page": 1, "lines": pages}
                    ]
                # Final append – this is what simple_keyword_search expects
                file_data.append({
                    "name": relative_name,     # was effectively blob.name before
                    "full_path": blob.name,
                    "content": content_string,  # SAME field your search uses
                    "pages1": pages              # NEW, but extra only
                })

            except Exception as e:
                print(f"Error processing {blob.name}: {e}")
                continue


        print(f"✅ Found {len(file_data)} usable documents in directory '{directory_path}'.")
        return file_data

    except Exception as e:
        print(f"Error listing/downloading files from GCS: {e}")
        return []


def highlight_matches_html(text: str, words: List[str], match_type: str = "partial") -> str:
    """Wraps matching words in HTML for highlighting (BiDi safe)."""
    if not words: return text

    all_patterns = []
    for w in words:
        if w:
            all_patterns.append(re.escape(w))
            all_patterns.append(re.escape(w[::-1]))

    if not all_patterns: return text

    unique_patterns = list(set(all_patterns))

    if match_type == "full":
        pattern = r"\b(" + "|".join(unique_patterns) + r")\b"
    else:
        pattern = r"(" + "|".join(unique_patterns) + r")"

    regex = re.compile(pattern, re.IGNORECASE)

    def repl(m):
        return f"<span style='background-color: #ffda79; color: #1f2937; font-weight:bold; padding: 2px 4px; border-radius: 4px;'>{m.group(0)}</span>"

    return regex.sub(repl, text)