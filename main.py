import sys, boto3
import email.utils
import os, io
import json
import requests
from PyQt5 import QtWidgets, QtCore
from docx import Document
from PIL import Image
import re
import ui_setup  # import your setup module
import win32com.client as win32
import pdfplumber
import unicodedata
from typing import List, Dict
import subprocess
from urllib.parse import urlparse, unquote
import mimetypes
import time
import threading
from google import genai
from google.genai.errors import APIError
from google.genai import types
import tempfile
from search_core import simple_keyword_search
from document_parsers import extract_text_and_images_from_pdf
from gcs_path_browser import GCSBrowserDialog, check_sync
from email_option_gui import launch_search_dialog
from email_searcher import EmailSearchWorker, EMAIL_PROVIDERS
from outlook_searcher import OutlookAPISearcher
from gmail_searcher import GmailAPISearcher
from icloud_searcher import ICloudAPISearcher
from speech2text import StopDialog
from utils import QRadioButton_STYLE_QSS_green_1515bg, QRadioButton_STYLE_QSS_green_1520bg, CHECKBOX_STYLE_QSS_green, CHECKBOX_STYLE_QSS_red
from search_utilities import initialize_all_clients
from ui_setup import non_sync_cloud_str, sync_cloud_str
import pytesseract
from utils import CHECKBOX_STYLE_QSS_black, CHECKBOX_STYLE_QSS_gray, Container_STYLE_QSSgray, Container_STYLE_QSS
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\Tesseract-OCR\tesseract.exe'

chat_gpt = False

api_key = os.environ.get("GEMINI_API_KEY")  # $Env:OPENAI_API_KEY = "sk-proj-vf-..."  echo $Env:OPENAI_API_KEY in power shell
MODEL_NAME = 'gemini-2.5-flash'
chat_mode = 'gemini'




RLO = u"\u202e"
CONFIG_FILE = "config.txt"
HEADER_MARGIN = 50
FOOTER_MARGIN = 50

chatgptallfile = True #False True
# Combined pattern for any character considered a 'letter' (Latin or Hebrew)
LETTER_PATTERN = re.compile(r'[a-zA-Z\u0590-\u05FF]')
# Define a pattern to find any Hebrew character
HEBREW_PATTERN = re.compile(r'[\u0590-\u05FF]')
ENGLISH_PATTERN = re.compile(r'[a-zA-Z]')

# Pre-compile the pattern for performance
LATIN_LETTER_PATTERN = re.compile(r'[a-zA-Z]')
sequence_pattern = re.compile(r'\d+')
LATIN_LETTER_PATTERNnNum = re.compile(r'[a-zA-Z]+|\d+')

from config_reader import BUCKET_NAME, API_search_url, API_simple_search_url, API_cache_status_url, CLIENT_PREFIX_TO_STRIP







# Function to implement in your class:
def display_keyword_matches(self, match_results):
    display_text = "<h3>Keyword Search Results</h3>"

    if not match_results:
        display_text += "<p>No matches found with the current configuration.</p>"

    for match in match_results:
        doc_name = match.get('document_name', 'Unknown Document')
        match_type = match.get('match_type', 'N/A')
        snippets = match.get('snippets', [])

        display_text += f"<hr><h4>File: {doc_name}</h4>"
        display_text += f"<p>Match Type: <b>{match_type}</b></p>"

        # Display up to the first 5 snippets for brevity
        for i, snippet in enumerate(snippets[:5]):
            display_text += f"<p>— Snippet {i + 1}: {snippet}</p>"

    return display_text

def format_simple_search_results(results_data):
    if results_data.get("status") != "ok":
        return f"🛑 שגיאה: {results_data}"

    matches = results_data.get("matches", [])
    if not matches:
        return "לא נמצאו תוצאות."

    output_lines = []

    for doc in matches:
        file_name = doc.get("file", "ללא שם")
        full_path = doc.get("full_path", "")
        match_positions = doc.get("match_positions","")
        if match_positions:
            first = match_positions[0]
            line = first["line"]
            page = first["page"]
        else:
            line = None
            page = None

        dir_only = os.path.dirname(full_path)

        lines = doc.get("matches_html", [])

        output_lines.append(f" שורה:  {line}  עמוד: {page}  📄 קובץ: {file_name}  📄 ספריה: {dir_only}   <br>")

        output_lines.append(f" debug:  {results_data.get("debug")}<br>***********************************************<br>")
        #output_lines.append(f"נתיב מלא: {full_path} <br>")


        for line in lines:
            output_lines.append(f"   • {line}<br>")

        output_lines.append(f"<br>")


    return "\n".join(output_lines)


def check_cache_status_get():
    print(f"\n--- 2. Checking Cache Status (GET {API_cache_status_url}) ---")

    # Use requests.get() for status checks


    try:
        while True:
            response = requests.get(API_cache_status_url)
            response.raise_for_status()
            status_data = response.json()
            print(f"Status Check Successful (Status {response.status_code}):")
            print(json.dumps(status_data, indent=4))
            print(f"\nCache Status: {status_data.get('status')} | Docs Cached: {status_data.get('document_count')}")
            cache_value = status_data['cache']
            process_value = status_data['process']
            print(f"process_value",process_value)
            print(f"cache_value", cache_value)

            #self.progressBar.setValue(int(process_value))
            if process_value > 99: #status_data.get('status') == 'WARMING_UP':
                break
            #    print("Waiting a moment for the background thread to finish...")
            #    self.progressBar.setValue(process_value)
            #    time.sleep(1)
            #    check_cache_status_get(self)  # Check again using GET
            time.sleep(1)

        return process_value

    except requests.exceptions.HTTPError as e:
        print(f"HTTP Error: {e}")
    except requests.exceptions.RequestException as e:
        print(f"Request Error: {e}")


def start_monitoring_job():
    # 1. Start the server-side job (must be non-blocking on the server side!)
    # ... (requests.post() to the /simple_search_key endpoint) ...

    # 2. Start the local polling thread
    poll_thread = threading.Thread(       target=check_cache_status_get,
        daemon=True  # Set to True so it automatically stops when the main app exits
    )
    poll_thread.start()

def on_search_button_clicked(self, query, directory_path ,force_chat = False):
    # הפונקציה המדויקת ששלחת, בתוספת לוגיקת עיבוד התוצאה.
    #self.search_button.setEnabled(False)
    #self.progressBar.setValue(0)  # Reset the progress bar
    try:
        #start_monitoring_job()
        # 1. שליחת הבקשה ל-Cloud Run API

        if self.gemini_radio.isChecked():
            if self.isLTR:
                query += " please indicate in which documents you found the answer "
            else:
                query += " - (פרט היכן מצאת את המידע) "
            if self.cloud_storage_provider == "Google":
                url = API_search_url
                payload = {
                    "query": query,
                    "directory_path": directory_path
                }
            elif self.cloud_storage_provider == "Amazon":
                AWS_REGION = "ap-southeast-2"
                KB_ID = "SEL4HSGWF3"
                # 1. Connect to AWS
                session = boto3.Session(
                    aws_access_key_id=os.environ.get("AWS_ACCESS_KEY"),
                    aws_secret_access_key=os.environ.get("AWS_SECRET_KEY"),
                    region_name=AWS_REGION
                )

                # 2. Use the 'bedrock-agent-runtime' client for Knowledge Bases
                client = session.client('bedrock-agent-runtime')

                # 3. Ask the AI to search S3 and generate an answer

                response = client.retrieve_and_generate(
                    input={'text': query},
                    retrieveAndGenerateConfiguration={
                        'type': 'KNOWLEDGE_BASE',
                        'knowledgeBaseConfiguration': {
                            'knowledgeBaseId': KB_ID,
                            'modelArn': f'arn:aws:bedrock:{AWS_REGION}::foundation-model/anthropic.claude-3-5-sonnet-20241022-v2:0',
                            'retrievalConfiguration': {
                                'vectorSearchConfiguration': {
                                    'numberOfResults': 10,  # Search 10 chunks instead of 5
                            #        'overrideSearchType': 'HYBRID'  # Uses Keywords + Meaning
                                }
                            }
                        }
                    }
                )

                return response['output']['text']
        else:

            if self.all_word_search_radio.isChecked():
                str2_mode = "all"
            else:
                str2_mode = "any"

            if self.exact_search_radio.isChecked():
                str1_mode = "full"
            else:
                str1_mode = "partial"

            if self.show_line_mode_radio.isChecked():
                str3_mode = "line"
            else:
                str3_mode = "paragraph"


            payload = {
                "query": query,
                "directory_path": directory_path,
                "search_config": {
                    "mode": "keyword",
                    "match_type": str1_mode,
                    "word_logic": str2_mode,
                    "show_mode": str3_mode
                }
            }

        start_time = time.time()
        if not self.cloud_gemini_radio.isChecked() and not force_chat:
            results_data = simple_keyword_search(query, directory_path, str1_mode, str2_mode, str3_mode)
            formatted_output = format_simple_search_results(results_data)
            return formatted_output
        else:
            if self.cloud_storage_provider == "Google":
                url = API_simple_search_url
                response = requests.post(
                    url,
                    json=payload,
                    headers={'Content-Type': 'application/json'}
                )
            elif self.cloud_storage_provider == "Amazon":
                a = 0
        # --- 2. מדידת זמן: התחלה ---

        end_time = time.time()
        latency = end_time - start_time
        print(f"requests.post: {latency:.2f} seconds")
        response.raise_for_status()  # לזרוק שגיאה אם הסטטוס הוא 4xx/5xx


        # 2. עיבוד התוצאה
        results_data = response.json()

        # 3. בודק אם הסטטוס הוא RAG (מ-Gemini) או Fallback (חיפוש פשוט)
        status = results_data.get('status', 'Unknown')

        formatted_output = ""

        if status.endswith('(RAG)'):
            # --- פורמט RAG (תשובה אנליטית מ-Gemini) ---
            response_text = results_data.get('response', 'הבינה המלאכותית לא סיפקה תשובה ברורה.')

            formatted_output = (
                f"✅ **תשובה אנליטית (Gemini AI)**\n"
                f"----------------------------------------\n"
                f"{response_text}\n\n"
            )

            # אם יש מקורות, ניתן להוסיף אותם כאן (כרגע הקוד לא מוציא אותם)
        elif status == "ok":

            # --- Keyword Search Mode ---
            formatted_output = format_simple_search_results(results_data)


        elif status.endswith('(Fallback)'):
            # --- פורמט Fallback (חיפוש מילות מפתח) ---

            # עדיף להציג הודעת שגיאה ברורה למה ה-RAG נכשל
            formatted_output = (
                f"⚠️ **כשל ב-RAG: המערכת נפלה לחיפוש מילות מפתח.**\n"
                f"----------------------------------------------------\n"
                f"הודעת מערכת: {results_data.get('message', 'לא נמצאו מסמכים.')}\n\n"
            )

            # הוספת הסניפטים מה-Fallback
            if results_data.get('results'):
                for result in results_data['results']:
                    formatted_output += (
                        f"📄 **{result['filename']}** (קטע רלוונטי):\n"
                        f"   {result['snippet']}\n"
                        f"----------------------------------------\n"
                    )

        else:
            # טיפול בשגיאות מה-API עצמו
            formatted_output = f"🛑 **שגיאת API או סטטוס לא ידוע:**\n{results_data.get('message', 'לא התקבלה הודעת שגיאה.')}"

        return formatted_output

    except requests.exceptions.RequestException as e:
        # טיפול בשגיאות חיבור לרשת או שגיאות HTTP
        error_message = f"שגיאת תקשורת עם שירות החיפוש (API):\n{e}"
        print(error_message)

        # הצגת הודעה קריטית למשתמש
        QtWidgets.QMessageBox.critical(self, "שגיאה", error_message)


def should_skip_file(filename: str) -> bool:
    """Checks if a filename indicates a temporary Word lock file."""

    # The standard temporary lock file prefix is '~$', but sometimes
    # the order might vary slightly due to OS or application behavior.

    # Check for the primary temporary file prefixes:
    if filename.startswith('~$'):
        return True

    # Check for the alternate temporary file prefix (less common but good practice)
    if filename.startswith('~$'):
        return True

    return False

def display_gemini_result(self, results_area: QtWidgets.QTextEdit, answer: str, path: str):
    """
    Handles displaying the Gemini answer in the results area.

    Appends an asterisk inline if the answer is a skip condition ('#$$$#' or 'None'),
    otherwise, appends the full formatted answer in a new RTL block.

    Args:
        self: The main class instance (needed for moveCursor/insertHtml).
        results_area: The QTextEdit widget where results are displayed.
        answer: The text response from the Gemini model.
        path: The file path used for context.
    """

    # Check for skip conditions: flag or explicit 'None' string
    if ('#$$$#' in answer) or answer == 'None':
        # --- Handle Skip Condition (Insert Asterisk Inline) ---

        # Use a <span> tag (inline element) with bolding
        html_content = f"<span dir='rtl'><b>*</b></span>"

        # Move cursor to the end and insert content without a newline
        results_area.moveCursor(results_area.textCursor().End)
        results_area.insertHtml(html_content)

    else:
        # --- Handle Successful Answer (Insert Full Block) ---

        # Use <p> tags for clear line separation, relying on append() for the initial block
        formatted_html = (
            f"<p dir='rtl'><b>    Gemini :  </b></p>"
            f"<p dir='rtl'><b>  {path}   </b></p>"
            f"<p dir='rtl'>{answer}</p>"
            f"<p dir='rtl'><b>    +++++++++++++++++++++ :  </b></p>"
        )
        # The append() method creates a new block for the formatted HTML content.
        results_area.append(formatted_html)



def extract_text_from_docx(docx_path: str) -> str:
    """Extracts all text from a local .docx file."""
    if not os.path.exists(docx_path):
        print(f"❌ Error: File not found at {docx_path}.")
        return None

    try:
        # Load the document
        document = Document(docx_path)
        text = []

        # Iterate through all paragraphs and append text
        for paragraph in document.paragraphs:
            text.append(paragraph.text)

        print(f"✅ Successfully extracted {len(' '.join(text).split())} words from the DOCX.")

        # Join all paragraph text into a single string
        return "\n".join(text)

    except Exception as e:
        print(f"❌ Error reading .docx file: {e}")
        return None


def extract_docx_content(docx_path: str, question: str):
    """
    Extracts text and embedded images from a DOCX file and prepares the Gemini contents list.
    """
    contents_list = []

    # 1. Create a temporary directory to save the extracted images
    with tempfile.TemporaryDirectory() as temp_dir:
        # 2. Extract text and save images to the temp directory
        try:
            matches = docx_search(docx_path, question)
            full_text = matches #docx2txt.process(docx_path, temp_dir)

        except Exception as e:
            print(f"Error during docx extraction: {e}")
            return [types.Part.from_text(f"Could not process document. Question: {question}")]

        # 3. Process and append images to the contents_list
        # FIX 1: Exclude Word lock files (starting with ~$)
        image_files = [
            f for f in os.listdir(temp_dir)
            if f.lower().endswith(('.png', '.jpg', '.jpeg'))
            and not f.startswith('~$')
        ]

        for image_file in sorted(image_files):
            image_path = os.path.join(temp_dir, image_file)

            # 🛑 FIX 2 (CRITICAL): Defensive check MUST be done BEFORE opening the file.
            # This handles the phantom second execution failure.
            if not os.path.exists(image_path):
                print(f"Skipping file {image_file}: File not found during execution.")
                continue

            # Use a try/except for *each individual image*
            try:
                with open(image_path, 'rb') as f:
                    image_bytes = f.read()

                    # Determine MIME type
                    mime_type = 'image/jpeg' if image_file.lower().endswith(('.jpg', '.jpeg')) else 'image/png'

                    # Add the image part
                    contents_list.append(
                        types.Part.from_bytes(data=image_bytes, mime_type=mime_type)
                    )

                    # Add the text marker
                    # FIX 3: Removed redundant 'text =' keyword for clarity.
                    contents_list.append(types.Part.from_text(text = f"--- Embedded Image: {image_file} ---"))

            except Exception as e:
                print(f"Failed to process image {image_file}: {e}")

        # --- END OF LOOP ---

        # 4. Define and append the final text prompt (only once, outside the loop)
        final_prompt = (
            f"--- EXTRACTED DOCX TEXT ---\n"
            f"{full_text}\n"
            f"--- INSTRUCTIONS ---\n"
            f"QUESTION: {question}"
        )
        a = types.Part.from_text(text = final_prompt)
        contents_list.append(a)

        # 5. Return the result (This must be the last line inside the 'with' block)
        return contents_list



def ask_gemini_with_context(context: str, question: str):
    """Asks Gemini a question using the PDF text as context."""
    if not context:
        return "Could not read the document to provide context."

    # 1. Construct the RAG Prompt
    # This prompt tells Gemini to use the provided text as its sole source of truth.
    system_instruction = (
        "You are a helpful assistant. Use ONLY the provided document text as context "
        "to answer the question. If the information is not in the text, reply #$$$# "
        #"the answer cannot be found in the provided document."
    )

    full_prompt = (
        f"DOCUMENT CONTEXT:\n---\n{context}\n---\n\n"
        f"QUESTION: {question}"
    )

    try:
        client = genai.Client()

        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=full_prompt,
            config={"system_instruction": system_instruction}
        )

        return response.text.strip()
        print("\n--- Gemini's Answer (Based on the PDF) ---")
        print(response.text.strip())
        print("------------------------------------------")

    except APIError as e:
        print(f"\n❌ API Error: Check your GEMINI_API_KEY and API access. Details: {e}")
    except Exception as e:
        print(f"\n❌ An unexpected error occurred: {e}")

def ask_gemini_with_contextNimage(context: str, image_list, question: str):
    """Asks Gemini a question using the PDF text as context."""
    if not context:
        return "Could not read the document to provide context."

    contents_for_gemini = []

    # Convert each PIL Image object into the required 'Part' format
    for image in image_list:
        # Use image.tobytes() if it's a PIL Image object
        # OR use the image bytes directly if your converter returns bytes
        # This example assumes the image object can be saved to bytes:
        #from io import BytesIO
        #img_byte_arr = BytesIO()
        #image.save(img_byte_arr, format='PNG')
        #img_bytes = img_byte_arr.getvalue()

        contents_for_gemini.append(
            types.Part.from_bytes(data=image, mime_type='image/png')
        )

    # 1. Construct the RAG Prompt
    # This prompt tells Gemini to use the provided text as its sole source of truth.

    system_instruction = (
        "You are a helpful assistant. Use ONLY the provided document text as context and the images"
        "to answer the question. If the information is not in the text, reply #$$$# "
        #"the answer cannot be found in the provided document."
    )

    final_prompt = (
        f"DOCUMENT CONTEXT:\n---\n{context}\n---\n\n"
        f"QUESTION: {question}"
    )

    contents_for_gemini.append(final_prompt)
    try:
        client = genai.Client()

        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=contents_for_gemini,
            config={"system_instruction": system_instruction}
        )

        return response.text.strip()
        print("\n--- Gemini's Answer (Based on the PDF) ---")
        print(response.text.strip())
        print("------------------------------------------")

    except APIError as e:
        print(f"\n❌ API Error: Check your GEMINI_API_KEY and API access. Details: {e}")
    except Exception as e:
        print(f"\n❌ An unexpected error occurred: {e}")

def open_pdf_page(pdf_path, page_number):
    sumatra_path = r"C:\Users\orenm\AppData\Local\SumatraPDF\SumatraPDF.exe"  # change this to your installation
    subprocess.Popen([sumatra_path, "-page", str(page_number), pdf_path])

def is_mostly_english(word: str, threshold: float = 0.75) -> bool:
    """
    Checks if a word is composed of at least 75% Latin letters (a-Z).
    """
    if not word:
        return False

    # 1. Find all Latin letters
    latin_count = len(LATIN_LETTER_PATTERN.findall(word))
    num_count = len(sequence_pattern.findall(word))

    # 2. Total length of the word
    total_chars = len(word)

    # 3. Calculate the ratio
    ratio = (latin_count+num_count) / total_chars

    return ratio >= threshold


def contains_hebrew(text: str) -> bool:
    """Checks if the string contains any Hebrew characters."""
    return bool(HEBREW_PATTERN.search(text))


def contains_english(text: str) -> bool:
    """Checks if the string contains any Hebrew characters."""
    return bool(ENGLISH_PATTERN.search(text))


def reverse_first_to_last_non_letter_span(word: str) -> str:
    """
    Finds the first and last non-letter characters in a word and reverses
    the entire substring between those two points (inclusive).
    """
    first_idx = -1
    last_idx = -1

    # 1. Find the index of the first non-letter character
    for i, char in enumerate(word):
        if not LETTER_PATTERN.match(char):
            first_idx = i
            break

    # 2. Find the index of the last non-letter character (iterate backward)
    for i in range(len(word) - 1, -1, -1):
        if not LETTER_PATTERN.match(word[i]):
            last_idx = i
            break

    # If no non-letters were found, return the original word
    if first_idx == -1:
        return word

    # 3. Define the three parts of the word: prefix, span, and suffix
    prefix = word[:first_idx]
    span_to_reverse = word[first_idx: last_idx + 1]
    suffix = word[last_idx + 1:]

    # 4. Reverse the span and reconstruct the word
    reversed_span = span_to_reverse[::-1]

    return prefix + reversed_span + suffix


def fix_hebrew_reversal(text):
    """
    Applies the advanced, surgical fix:
    1. Always reverses characters within each word.
    2. Then, surgically reverses non-letter segments (numbers/punctuation)
       back to LTR order to fix corruption.
    3. Finally, reverses the word order in the sentence.
    """
    if not text:
        return ""

    def is_hebrew_letter(char):
        # Checks if a character is a standard Hebrew letter
        return 'HEBREW' in unicodedata.name(char, '')

        corrected_segments = []
        for non_letter_segment, letter_segment in segments:
            if letter_segment:
                # If it's a letter segment (Hebrew or Latin), keep it as-is (it's already reversed)
                corrected_segments.append(letter_segment)
            elif non_letter_segment:
                # If it's a non-letter segment (numbers, punctuation), reverse it back
                # Example: '321' needs to be reversed back to '123'
                corrected_segments.append(non_letter_segment[::-1])

        return "".join(corrected_segments)

    # 1. Split the text into words by whitespace
    words = text.split()

    surgically_corrected_words = []
    for word in words:
        # Step 1: Always apply full character reversal first

        # if ~contains_english(word):
        #    fully_reversed_word = word[::-1]
        # else:
        #    fully_reversed_word = word
        if is_mostly_english(word):
            fixed_word = word
        else:
            fully_reversed_word = word[::-1]

            fully_reversed_word = fully_reversed_word.replace(')', 'zzzz').replace('(', 'xxxx').replace('"', 'cccc').replace('.', 'dddd')
            #fully_reversed_word = fully_reversed_word.replace('"', 'cccc')
            fixed_word = reverse_first_to_last_non_letter_span(fully_reversed_word)
            #fixed_word = fixed_word.replace('cccc', '"')
            fixed_word = fixed_word.replace('zzzz', ')').replace('xxxx', '(').replace('cccc', '"').replace('dddd', '.')

        surgically_corrected_words.append(fixed_word)

    # 3. Reverse the order of the words in the list (fixes RTL sentence order)
    surgically_corrected_words.reverse()

    # 4. Join the words back with spaces
    return " ".join(surgically_corrected_words)


def paragraph_matches(text, words, mode='any', search_mode='partial'):
    # text_lower = text.lower()
    text_lower = text
    if search_mode == 'full':
        # Match only whole words
        found = 0
        for w in words:
            pattern = r'\b{}\b'.format(re.escape(w.lower()))
            if re.search(pattern, text_lower):
                found += 1

        return (mode == 'all' and found == len(words)) or (mode != 'all' and found > 0)
    else:
        # substring match
        if mode == 'all':
            return all(w.lower() in text_lower for w in words)
        else:
            return any(w.lower() in text_lower for w in words)


def replace_with_bold(text, regex):
    # Use QRegExp to find matches and replace with <b> tags
    pattern = regex
    start = 0
    result = ""
    while True:
        index = pattern.indexIn(text, start)
        if index == -1:
            result += text[start:]
            break
        # Append text before match
        result += text[start:index]
        # Append bolded match
        match_text = pattern.cap(0)
        result += f"<b>{match_text}</b>"
        start = index + len(match_text)
    return result


def highlight_terms(text, search_terms):
    # Highlight each search term in the text (case-insensitive)
    for term in search_terms:
        # Use regex for case-insensitive replacement
        pattern = QtCore.QRegExp(term, QtCore.Qt.CaseInsensitive)
        text = replace_with_bold(text, pattern)
    return text


def convert_docx2pdf(docx_path, pdf_path):
    word = win32.Dispatch('Word.Application')
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        # doc = Document(os.path.abspath(docx_path))
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
    except Exception as e:
        print("Error:", e)
    finally:
        word.Quit()


def docx2pdf_search(path, words, mode='any', search_mode='partial'):
    pdf_path = os.path.join(os.getcwd(), "tempXX.pdf")
    convert_docx2pdf(path, pdf_path)
    results = pdf_search( path, words, mode, search_mode, pdf_path)
    # Optionally delete the temp.pdf here if not done inside pdf_search
    try:
        os.remove(pdf_path)
    except:
        pass
    return results


def pdf_search(self, path, words, mode='any', search_mode='partial', read_from_temp=""):
    results = []

    try:
        if read_from_temp.strip():
            pdf_path = read_from_temp
            if os.path.exists(pdf_path):
                reader = pdfplumber.open(pdf_path)
            else:
                print("Temporary PDF not found.")
                return results
        else:
            reader = pdfplumber.open(path)

        for pnum, page in enumerate(reader.pages, start=1):

            # 1. ATTEMPT TO EXTRACT STRUCTURED TEXT (Current Logic)
            page_height = page.height
            header_limit = HEADER_MARGIN
            footer_limit = page_height - FOOTER_MARGIN

            blocks: List[Dict] = page.extract_text_lines(
                return_chars=False,
                return_textbox=True,
                strip=False
            )

            # Apply filtering and sorting logic as before
            filtered_blocks = [
                block for block in blocks
                if block['top'] > header_limit and block['bottom'] < footer_limit
            ]
            sorted_blocks = sorted(filtered_blocks, key=lambda b: (b['top'], b['x0']))
            page_raw_text_lines = [block['text'] for block in sorted_blocks]

            # Initialize lines from structured text
            lines = [fix_hebrew_reversal(line) for line in page_raw_text_lines]

            # 2. OCR FALLBACK LOGIC
            # If pdfplumber found very few lines (e.g., only page numbers or nothing), run OCR
            if len("".join(lines).strip()) < 50:  # Check if structured text is sparse

                print(f"Page {pnum}: Structured text is sparse. Running OCR fallback.")

                try:
                    # Render the page as a PIL image
                    # Note: You may need to install 'poppler' utilities for pdfplumber to work with image rendering
                    rgb_image_object = page.to_image(resolution=220).original
                    im = rgb_image_object.convert('L')
                    if self.isLTR:
                        ocr_text = pytesseract.image_to_string(im, lang='eng')
                    else:
                        ocr_text = pytesseract.image_to_string(im, lang='heb')

                    lines = [line.strip() for line in ocr_text.split('\n') if line.strip()]


                    if not lines:
                        print(f"Page {pnum}: OCR failed to find text.")

                except Exception as e:
                    print(f"Page {pnum}: OCR failed or Tesseract is not configured: {e}")

            # --- CONTINUE WITH SEARCH LOGIC (The rest of your function remains mostly the same) ---

            previous_page_trailer = []
            lines = previous_page_trailer + lines  # Re-initialize lines here after OCR check

            l = len(lines)

            i = 0
            if search_mode == 'chatgpt':
                # Existing ChatGPT logic...
                content_summary = "\n".join(lines)

                query = self.search_input.toPlainText().strip()
                prompt = f"כתוב את התשובה בשפה העברית על בסיס הטקסט הבא רק אם יש לך תשובה החלטית אחרת השב $$$$:\n{content_summary}\nשאלה: {query}"
                pre = f"<span style='color:blue;'>— עמוד {pnum}</span>"
                time.sleep(2)
                answer = self.ask_chatgpt(prompt)
                pre = f"<span style='color:blue;'> — עמוד {pnum}  </span>"

                full_paragraph = (
                        " ".join([path]) + "  " +
                        pre + "<br><br>" +
                        answer + "<br>"
                )

                if "error" in answer.lower() or "אין מספיק מידע" in answer.lower() or "$$$$" in answer.lower() or "לא מספק מידע" in answer.lower() or "אני מצטער" in answer.lower() or "בטקסט שסופק" in answer.lower():
                    full_paragraph = ""

                results.append(full_paragraph)

            else:
                while i < l:
                    lnum = i + 1
                    ln = lines[i]

                    # 1. First Check: Does the current line match ANY of the search words?
                    if paragraph_matches(ln, words, 'any', search_mode):
                        # 2. Safely gather the 3-line context
                        start_index = max(0, lnum - 2)
                        end_index = min(lnum + 2, l)
                        if (lnum + 2) > l:
                            previous_page_trailer = lines[lnum - 1:]
                        else:
                            previous_page_trailer = []

                        context_lines = lines[start_index:end_index]
                        context_lines = [line.replace(".₪", "₪.") for line in context_lines]
                        context_lines = [line.replace(",₪", "₪,") for line in context_lines]

                        # 3. Concatenate lines into a single searchable string
                        context_text = " ".join(context_lines)

                        # 4. Second Check: Does the 3-line context match the FULL user criteria?
                        if paragraph_matches(context_text, words, mode, search_mode):
                            pre = f"<span style='color:blue;'>— עמוד {pnum} — שורות {start_index + 1}-{end_index}</span>"
                            path_for_url = path.replace('\\', '/')
                            file_url_with_page = f"filepage:///{path_for_url}?page={pnum}"
                            open_link = f"<a href='{file_url_with_page}' style='color:green; text-decoration: none;'>[פתח קובץ]</a>"

                            full_paragraph = (
                                    " ".join([path]) + "  " +
                                    pre + " " + open_link + "<br><br>" +
                                    "<br>".join(context_lines) + "<br>"
                            )

                            results.append(full_paragraph)
                            i += 2

                    i = i + 1

    except Exception as general_error:
        print(f"An unexpected error occurred during PDF processing: {general_error}")

    finally:
        # Delete temp PDF if it was used
        if read_from_temp and os.path.exists(pdf_path):
            try:
                os.remove(pdf_path)
            except Exception as e:
                print(f"Error deleting temp PDF: {e}")

    return results



def docx_search(self, path, words, mode='any', search_mode='partial'):
    results = []
    try:
        doc = Document(path)
        for para in doc.paragraphs:
            full_text = para.text or ""
            if paragraph_matches(full_text, words, mode, search_mode):
                #results = docx2pdf_search(path, words, mode, search_mode)
                #return results

                full_paragraph = (
                    f"{path}  <br><br> {full_text} <br>"
                )
                results.append(full_paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text = cell.text or ""
                    full_paragraph = (
                            " ".join([path]) + "  " +
                            "<br><br>" +
                            "<br>".join(full_text) + "<br>"
                    )

                    if paragraph_matches(full_text, words, mode, search_mode):
                        results.append(full_paragraph)

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                # 1. Extract Image Data
                image_part = rel.target_part
                image_bytes = image_part.blob

                # 2. Open Image using Pillow
                image = Image.open(io.BytesIO(image_bytes))

                # 3. Run OCR on the Image
                try:
                    # Use Tesseract to convert image to string
                    image = image.convert('L')

                    if self.isLTR:
                        ocr_text = pytesseract.image_to_string(image, lang='eng')
                    else:
                        ocr_text = pytesseract.image_to_string(image, lang='heb')

                    if paragraph_matches(ocr_text, words, mode, search_mode):
                        full_paragraph = (
                        f"{path}  <br><br> {ocr_text} <br>"
                        )
                        results.append(full_paragraph)
                except pytesseract.TesseractNotFoundError:
                    print("🚨 Tesseract is not installed or the path is incorrect. Cannot perform OCR.")
                    return None
                except Exception as e:
                    print(f"An error occurred during OCR: {e}")

    except:
        pass

    return results




class SearchApp(QtWidgets.QWidget):
    def set_window_title(self):
        if not self.cloud_gemini_radio.isChecked():
            self.setWindowTitle(f"הדס לוי -  עורך דין - תוכנת חיפוש " + f" Hard Disk")
        else:
            if self.cloud_storage_provider == "Google":
                self.setWindowTitle(
                    f"  הדס לוי -  עורך דין - תוכנת חיפוש  {self.cloud_run_rev} -  {self.cloud_storage_provider} ")
            elif self.cloud_storage_provider == "Amazon":
                self.setWindowTitle(f"  הדס לוי -  עורך דין - תוכנת חיפוש  {self.cloud_storage_provider}")
            else:
                self.setWindowTitle(f"  הדס לוי -  עורך דין - תוכנת חיפוש  {self.cloud_storage_provider}")

    def __init__(self):
        super().__init__()
        self.email_worker = None
        self.thread = None

        ui_setup.setup_ui(self)
        if self.update_app_title:
            response = requests.get(API_cache_status_url)
            response_str = response.content.decode('utf-8').strip()
            data = json.loads(response_str)
            self.cloud_run_rev=data["REVISION"]

            self.set_window_title()


        self.resize(1800, 1200)

        result = check_sync(CLIENT_PREFIX_TO_STRIP+"/גירושין/", BUCKET_NAME, prefix='גירושין')
        sync0 = result["sync!"]
        self.sync0 = sync0
        self.update_gcs_radio()


    def speech2text_handler(self):
        # 1. Determine language and mode from your radio buttons
        # Replace 'radio_manual' with your actual radio button object name
        self.current_voice_result = ""
        lang = "he-IL" if self.Language == "Hebrew" else "en-US"

        # 2. Initialize the worker

        text = StopDialog.get_voice_text(parent=self, language=lang, mode=self.Voice_recognition_mode) # manual auto

        self.search_input.setText(text)


    def handle_chatgpt_mode(self, folder, query):
        # Prepare content from documents
        content_summary = self.prepare_content_summary(folder)

        # Create prompt
        #prompt = f"Based on the following content:\n{content_summary}\nQuestion: {query}"
        prompt = f"כתוב את התשובה בשפה העברית על בסיס הטקסט הבא ואל תשכח לציין את מקור המידע: שם הקובץ, ומספר העמוד:\n{content_summary}\nשאלה: {query}"

        # Call GPT API
        answer = self.ask_chatgpt(prompt)
        if "error" in answer.lower() or "אין מספיק מידע" in answer.lower() or "אני מצטער" in answer.lower()  or "בטקסט שסופק" in answer.lower()   :
            answer = ""
        return answer




    def update_search_button_text(self):
        layout = self.both_groups_layout  # Use the class attribute reference
        if self.gemini_radio.isChecked():
            self.search_btn.setText(ui_setup.gemini_radio_str)
            self.label.setText(ui_setup.label_gpt_str)
            self.g_group_widget.setVisible(False)

            self.nongemini_radio.setStyleSheet(QRadioButton_STYLE_QSS_green_1515bg)
            self.gemini_radio.setStyleSheet(QRadioButton_STYLE_QSS_green_1520bg)
        else:
            self.search_btn.setText(ui_setup.search_btn_str)
            self.label.setText(ui_setup.label_str)
            self.g_group_widget.setVisible(True)

            self.nongemini_radio.setStyleSheet(QRadioButton_STYLE_QSS_green_1520bg)
            self.gemini_radio.setStyleSheet(QRadioButton_STYLE_QSS_green_1515bg)

    # --- FINAL ATTEMPT: ShellExecuteW to bypass shell ambiguity ---
        # In SearchApp class
    def _handle_link_click(self, url: QtCore.QUrl):
        """
        Uses the Windows API function ShellExecuteW to reliably open the file,
        avoiding Explorer by forcing the native 'open' action on the file path.
        """
        # 1. Decode the local file path from the QUrl object
        # QUrl.toLocalFile() correctly handles URL-encoding and provides the clean
        # file path, including Hebrew characters.
        local_path = url.toLocalFile()
        url_str = url.toString()

        # Parse the URL to extract the path
        parsed_url = urlparse(url_str)
        # Get the path without URL encoding (e.g., spaces, non-ASCII)
        file_path = unquote(parsed_url.path)



        # Check for your custom protocol signal
        if url_str.startswith("filepage://"):
            try:
                # 2. Convert to absolute path (if needed) and use the correctly decoded path
                # os.path.abspath is usually good, but local_path from QUrl should be sufficient.
                # For Windows, on some Python versions, the path may have a leading slash
                if os.name == 'nt' and file_path.startswith('/'):
                    file_path = file_path[1:]

                # Now open the file

                content = url_str[len("filepage:///"):]  # everything after the scheme
                # Split to separate path and query
                path_part, _, query = content.partition('?')
                # Extract page number from query
                params = dict(param.split('=') for param in query.split('&'))
                page_num = int(params.get('page', 1))
                # Convert path back to OS format
                file_path = path_part.replace('/', os.path.sep)

                mime_type, _ = mimetypes.guess_type(file_path)

                if mime_type == 'application/pdf':
                    # Open PDF at specific page with SumatraPDF
                    open_pdf_page(file_path, page_num)
                else:
                    # Open other files normally
                    os.startfile(file_path)

            except Exception as e:
                print(f"Error launching file {file_path}: {e}")
                QtWidgets.QMessageBox.warning(
                    self, "שגיאה בפתיחת קובץ",
                    f"לא ניתן לפתוח את הקובץ: {file_path}\nשגיאה: {e}"
                )



    def clear_all(self):
        """Clears the search input and, most importantly, the results display area."""
        self.search_input.clear()
        self.results_area.clear()
        self.progressBar.setValue(0)

    def save_all2file(self):
        """Clears the search input and, most importantly, the results display area."""
        doc = Document()
        path = os.path.join(CLIENT_PREFIX_TO_STRIP, "results" + self.last_queryNmode + ".docx")
        text = self.results_area.toPlainText()
        # Add the text as a paragraph
        doc.add_paragraph(text)
        doc.save(path)




    def execute_search(self):
        query = self.search_input.toPlainText().strip()
        if self.cloud_storage_provider == "Google":
            if not initialize_all_clients():
                print("LOG: Request failed - Service initialization failed. Check server logs for IAM/API Key errors.")

        if not query:
            return

        if self.gemini_radio.isChecked():
            self.last_queryNmode = "_Chat_" + query
        else:
            self.last_queryNmode = "_Search_" + query

        start_time = time.time()
        print(f"Start search")
        self.search_btn.setText(ui_setup.press_search_btn_str)
        self.search_btn.setStyleSheet(CHECKBOX_STYLE_QSS_red)
        QtWidgets.QApplication.processEvents()
        self.search_btn.setEnabled(False)

        folders = self.dir_edit.text().strip()
        folder_list = [f.strip() for f in folders.split(',') if f.strip()]
        self.results_area.clear()
        total_found = 0

        for folder in folder_list:

            if self.cloud_gemini_radio.isChecked() or (self.non_cloud_gemini_radio.isChecked() and self.gemini_radio.isChecked() and self.hd_cloud_auto_toggle == "True"):

                if not self.sync0 or self.gemini_radio.isChecked() or not self.hd_cloud_auto_toggle:
                    if  self.non_cloud_gemini_radio.isChecked():
                        self.cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_black)
                        self.non_cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_gray)
                        self.display_root.setText("☁️ Bucket")
                        self.display_root.setStyleSheet("color: white; background-color: lightblue;")
                        QtWidgets.QApplication.processEvents()

                    normalized_path = folder.replace("\\", "/")
                    prefix_to_strip = CLIENT_PREFIX_TO_STRIP.replace("\\", "/")
                    # 3. Strip trailing slashes from both for consistent comparison (e.g., C:/a/מצרפי/)
                    normalized_path = normalized_path.strip('/')
                    prefix_to_strip = prefix_to_strip.strip('/')

                    # 4. Perform the strip only if the path starts with the prefix
                    if normalized_path.lower().startswith(prefix_to_strip.lower()):
                         gcs_directory_path = normalized_path[len(prefix_to_strip):].strip('/')
                    else:
                        # If no prefix match, use the normalized path as is
                        gcs_directory_path = normalized_path.strip('/')

                    folder_in = gcs_directory_path

                    end_time = time.time()
                    latency = end_time - start_time
                    print(f"pre on_search_button_clicked: {latency:.2f} seconds")

                    answer = on_search_button_clicked(self, query, folder_in, self.non_cloud_gemini_radio.isChecked())


                    #perform_search(query, directory_path=folder)
                    display_gemini_result(self, self.results_area, answer, folder_in)

                    if self.non_cloud_gemini_radio.isChecked():
                        self.non_cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_black)
                        self.cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_gray)
                        self.display_root.setText(CLIENT_PREFIX_TO_STRIP)


                    continue

                else:
                    self.non_cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_black)
                    self.cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_gray)
                    self.display_root.setText(CLIENT_PREFIX_TO_STRIP)




            folder_in = self.display_root.text() +  "/"+ folder
            if not os.path.isdir(folder_in):
                QtWidgets.QMessageBox.warning(self, "שגיאה", "התיקיה לא קיימת.")
                return
            if not query:
                QtWidgets.QMessageBox.warning(self, "שגיאה", "אנא הזן מילות חיפוש.")
                return

            # Parse query for AND/OR
            mode = 'any'
            # Determine search mode based on radio button selection
            if self.all_word_search_radio.isChecked():
                mode = 'all'


                # Determine mode
            if self.gemini_radio.isChecked():
                search_mode = chat_mode
            elif self.exact_search_radio.isChecked():
                search_mode = 'full'
            else:
                search_mode = 'partial'  # default


            words = [w.strip() for w in query.replace('+', ' ').replace('&', ' ').split()]

            count = 0
            for root, _, files in os.walk(folder_in):
                total_files = sum([len(files) for r, d, files in os.walk(folder_in)])
                for filename in files:
                    count += 1
                    if filename.startswith('~$'):
                        continue  # Skip temporary backup files

                    if should_skip_file(filename):
                        print(f"SKIPPING: {filename} (Temporary Lock File)")
                        continue


                    self.progressBar.setValue(int(count / total_files * 100))

                    QtWidgets.QApplication.processEvents()
                    path = os.path.join(root, filename)

                    path_str = f"{root} \t {filename}"
                    filename_lower = filename.lower()
                    if filename_lower.endswith('.docx'):
                        try:
                            if search_mode == 'gemini':

                                #doc_text = extract_text_from_docx(path)
                                #if doc_text:
                                #    answer = ask_gemini_with_context(doc_text, query)

                                gemini_input_parts = extract_docx_content(str(path), query)
                                client = genai.Client()
                                response = client.models.generate_content(
                                     model='gemini-2.5-flash',
                                     contents=gemini_input_parts
                                )
                                answer = response.text.strip()




                            # Example usage (assuming this function is part of a class):
                                display_gemini_result(self, self.results_area, answer, str(path))

                            else:
                                matches = docx_search(self, path, words, mode, search_mode)

                                for paragraph in matches:
                                    self.append_result(path, paragraph, words)
                                    total_found += 1
                        except:
                            continue
                    elif filename_lower.endswith('.pdf'):
                        try:

                            matches =""
                            if search_mode == 'gemini':
                                answer = ""
                                pdf_text, image_list = extract_text_and_images_from_pdf(str(path))
                                if image_list:
                                    answer = ask_gemini_with_contextNimage(pdf_text, image_list, query)
                                elif pdf_text:
                                    answer = ask_gemini_with_context(pdf_text, query)

                                display_gemini_result(self, self.results_area, answer, str(path))


                            else:
                                pdf_text, image_list = extract_text_and_images_from_pdf(str(path))
                                matches = pdf_search(self, path, words, mode, search_mode)

                            # matches2 = pdf_search_flags(path, words, search_mode)


                            for paragraph in matches:
                                self.append_result2(path, paragraph, words, search_mode )
                                total_found += 1
                        except:
                            continue

        self.search_btn.setText(ui_setup.search_btn_str)
        self.search_btn.setEnabled(True)
        self.search_btn.setStyleSheet(CHECKBOX_STYLE_QSS_green)

        if self.cloud_gemini_radio.isChecked() and self.sync0:
            self.cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_black)
            self.non_cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_gray)
            self.display_root.setText("☁️ Bucket")
            self.display_root.setStyleSheet("color: white; background-color: lightblue;")

        if  not self.gemini_radio.isChecked():
            if total_found == 0:
                self.results_area.append("לא נמצאו תוצאות.")
            else:
                self.results_area.append(f"\nנמצאו {total_found} תוצאות.")

    def save_last_dir(self):
        try:
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                f.write(self.dir_edit.text().strip())
        except:
            pass

    def display_email_results(self, results):
        # This method runs in the main GUI thread and receives 'results' list
        BLUE_STYLE = 'style="color: blue; font-weight: bold;"'
        self.g31_container.setStyleSheet(Container_STYLE_QSSgray)

        if not results:
            return

        email_type = results[0]
        search_params = self.email_search_params

        query = search_params["query"]
        directory = search_params["directory"]
        date_from_ts = search_params["date_from_ts"]
        date_to_ts = search_params["date_to_ts"] # add 24*3600sec hours to the end of day
        from_address = search_params["fromto_address"].lower()
        self.results_area.append(f"\n--- Email Search Results --- <span {BLUE_STYLE}>{query}</span> in <span {BLUE_STYLE}>{directory}</span>")
        email_address = email_name = ""

        self.last_queryNmode = "_Email_" + query

        for result in results[1:]:
            email_date_string = None
            email_timestamp = 0

            # --- 1. Date Extraction and Conversion ---
            lines = result.splitlines()
            email_from = ""
            # Iterate through the lines of the current email block (result)
            for line in lines:
                if line.startswith('Date:'):
                    # Extract the date string
                    parts = line.split(':', 1)
                    email_date_string = parts[1].strip()

                    # Convert to timestamp
                    time_tuple_tz = email.utils.parsedate_tz(email_date_string)
                    if time_tuple_tz:
                        email_timestamp = email.utils.mktime_tz(time_tuple_tz)

                elif line.startswith('From:'):
                    parts = line.split(':', 1)
                    email_from = parts[1].strip()
                    email_name, email_address = email.utils.parseaddr(email_from)
                    # --- 3. Extract Subject ---
                elif line.startswith('Subject:'):
                    parts = line.split(':', 1)
                    email_subject = parts[1].strip()

                    # --- 2. Filtering Logic ---

            # If the date couldn't be parsed, we might want to skip or include it.
            # Assuming we skip if the date is invalid (email_timestamp == 0)

            # Check the range: email_timestamp must be BETWEEN the from and to dates.
            # We use inclusive checks for the boundary dates.
            Time_Flag = date_from_ts <= email_timestamp <= date_to_ts
            From_Flag = not from_address or from_address in email_address.lower() or email_name == from_address

            if Time_Flag and From_Flag:
                # 3. Success: The email is in range, so append it.
                self.results_area.append(result)
            else:
                # 4. Failure: The email is outside the range.
                # Use 'continue' to skip the rest of this 'result' block and move to the next one.
                continue  # Skip to the next email block

    def display_gmail_results(self, results):
        search_params = self.email_search_params
        query = search_params["query"]
        self.last_queryNmode = "_" + search_params["provider_key"] + "_" + query
        directory = search_params["directory"]
        elapsed_time = time.time() - self.time0
        self.results_area.append(f"<b>Search completed in: {elapsed_time:.2f} seconds</b>")
        # This method runs in the main GUI thread and receives 'results' list
        BLUE_STYLE = 'style="color: blue; font-weight: bold;"'
        self.g31_container.setStyleSheet(Container_STYLE_QSSgray)

        if not results:
            self.results_area.append(
                f"\n--- Email Search Not Found !!! --- <span {BLUE_STYLE}>{query}</span> in <span {BLUE_STYLE}>{search_params["provider_key"]} {directory}</span>")
            return

        self.results_area.append(f"\n--- Email Search Results --- <span {BLUE_STYLE}>{query}</span> in <span {BLUE_STYLE}>{search_params["provider_key"]} {directory}</span>")




        for result in results:
            self.results_area.append(result)





    def email_search(self):
        self.g31_container.setStyleSheet(Container_STYLE_QSS)
        #QtWidgets.QApplication.processEvents()

        params=launch_search_dialog()

        self.email_search_params = params

        if not params:
            self.display_email_results("")
            return

        query = params["query"]
        folder = params["directory"]
        gmail_raw_query = params["gmail_raw_query"]

        fromto_address = params["fromto_address"]
        has_attachment = params["has_attachment"]
        date_from_ts = params["date_from_ts"]
        date_to_ts = params["date_to_ts"]
        min_size_kb = params["min_size_kb"]




        # --- ASSUME THESE ARE READ FROM NEW GUI INPUTS OR A CONFIG FILE ---
        email_user = params["email"]


        # Use re.findall() to extract all matches (the text within the capturing group)

        provider_key = params["provider_key"]

        if provider_key == "Gmail":
            email_password = "netj diso xxfv syqi"
        elif provider_key == "iCloud":
            email_password = "gjkk-momw-arkr-uhhv" #"Lael0404"
        else:
            email_password = "Jmjmjm2004"

        # Use the centralized dictionary to get server info
        provider_info = EMAIL_PROVIDERS.get(provider_key, {})
        server = provider_info.get("server")
        port = provider_info.get("port")

        if not server or not port:
            # Handle case where provider is not recognized
            self.results_area.append("ERROR: Selected email provider settings not found.")
            return

        # --- Instantiation with all parameters ---



        if provider_key == "Gmail":
            self.email_worker = GmailAPISearcher()
        elif provider_key == "Outlook":
            self.email_worker = OutlookAPISearcher()
        elif provider_key == "iCloud":
            self.email_worker = ICloudAPISearcher(email_user,email_password, server)
        else:
            self.email_worker = EmailSearchWorker(
                query,
                folder,
                email_user,
                email_password,
                server,
                port,
                params,
                provider_key
            )

        self.thread = QtCore.QThread()

        # 2. Move worker to the new thread
        self.email_worker.moveToThread(self.thread)

        # 3. Connect signals:
        #    a) When the thread starts, execute the worker's run() method.
        self.time0 = time.time()
        if provider_key == "Gmail":
            self.thread.started.connect(lambda: self.email_worker.search_emails_api(gmail_raw_query))
            self.email_worker.search_finished.connect(self.display_gmail_results)
        elif provider_key == "Outlook":
            self.thread.started.connect(lambda: self.email_worker.search_emails_api(query,fromto_address,has_attachment,min_size_kb,date_to_ts,date_from_ts))
            self.email_worker.search_finished.connect(self.display_gmail_results)
        elif provider_key == "iCloud":
            self.thread.started.connect(lambda: self.email_worker.search_emails_api(query, fromto_address, has_attachment, min_size_kb,date_to_ts, date_from_ts))
            self.email_worker.search_finished.connect(self.display_gmail_results)
        else:
            self.thread.started.connect(self.email_worker.search_emails_api)
            self.email_worker.search_finished.connect(self.display_email_results)
        #    b) When the worker finishes (emits the signal), process results and clean up.

        self.email_worker.search_finished.connect(self.thread.quit)  # Stops the thread loop

        #    c) Optional cleanup: delete worker and thread objects when thread finishes
        self.thread.finished.connect(self.email_worker.deleteLater)
        self.thread.finished.connect(self.thread.deleteLater)

        # 4. Start the thread (This runs worker.run() safely in the background)
        self.thread.start()

    from PyQt5 import QtWidgets, QtCore

    def select_files_and_folders(self, initial_path="/"):
        dialog = QtWidgets.QFileDialog(self)
        start_path = initial_path.replace("\\", "/")
        if os.path.exists(start_path):
            dialog.setDirectory(start_path)
        else:
            dialog.setDirectory(QtCore.QDir.homePath())

        dialog.setOption(QtWidgets.QFileDialog.DontUseNativeDialog, True)

        # 1. Set to Directory mode - this makes folders selectable
        # but normally limits you to one. We fix that next.
        dialog.setFileMode(QtWidgets.QFileDialog.Directory)

        # 2. Find the internal view (the list of files) and force 'Extended Selection'
        # This is what allows CTRL+Click to work for multiple folders
        for view in dialog.findChildren(QtWidgets.QAbstractItemView):
            view.setSelectionMode(QtWidgets.QAbstractItemView.ExtendedSelection)

        if dialog.exec_() == QtWidgets.QDialog.Accepted:
            # returns the list of all highlighted items (files or folders)
            return dialog.selectedFiles()
        return []

    def browse_directory(self):
        """
        Handles browsing local HD or GCS based on LOCAL_MODE.

        NOTE: 'self' refers to the main application window/widget, allowing it
        to pass itself as a parent to the dialog.
        """
        # Assuming self.dir_edit is a QLineEdit widget used to display the path


        if not self.cloud_gemini_radio.isChecked():

            root_dir = self.display_root.text()
            initial_path_from_gui = self.dir_edit.text()
            if "," in initial_path_from_gui:
                # Split by comma and take the first part, then strip extra spaces
                initial_path_from_gui = initial_path_from_gui.split(",")[0].strip()

            initial_path = root_dir + "/" + initial_path_from_gui



            dir_path_list = self.select_files_and_folders(initial_path)
            dir_path_str = ""
            # dir_path = RLO + dir_path # Assuming RLO is defined elsewhere
            if dir_path_list:
                prefix_to_strip = CLIENT_PREFIX_TO_STRIP.replace("\\", "/")
                prefix_to_strip = prefix_to_strip.strip('/')
                for dir_path in dir_path_list:
                # 4. Perform the strip only if the path starts with the prefix
                    if dir_path.lower().startswith(prefix_to_strip.lower()):
                        # Strip the prefix, plus the slash that separates the prefix from the folder path
                        dir_path = dir_path[len(prefix_to_strip):].strip('/')
                    else:
                        # If no prefix match, use the normalized path as is
                        dir_path = dir_path.strip('/')

                    dir_path_str = dir_path_str  + dir_path + ","
            self.dir_edit.setText(dir_path_str)
            self.save_last_dir()
        else:
            # --- GCS BROWSER INTEGRATION ---

            # 1. Use the static method to show the custom dialog
            initial_path_from_gui = self.dir_edit.text()
            if "," in initial_path_from_gui:
                # Split by comma and take the first part, then strip extra spaces
                initial_path_from_gui = initial_path_from_gui.split(",")[0].strip()

            initial_path = initial_path_from_gui



            selected_gcs_path = GCSBrowserDialog.get_directory(parent=self, initial_path=initial_path)

            # 2. Update the main application's path if a selection was made
            if selected_gcs_path is not None:
                # We add the bucket name conceptually, or just the path for internal use
                self.dir_edit.setText(selected_gcs_path)
                self.save_last_dir()

            # The previous debug printing block has been removed as it is now handled
            # by the interactive GCSBrowserDialog

    def load_last_dir(self):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                last_dir = f.read().strip()
                self.dir_edit.setText(last_dir)
        except:
            pass

    def update_gcs_radio(self):
        if self.sync0:
            self.cloud_gemini_radio.setText(sync_cloud_str)
            self.display_root.setStyleSheet("color: white; background-color: black;")
        else:
            self.cloud_gemini_radio.setText(non_sync_cloud_str)
            self.display_root.setStyleSheet("color: red; background-color: black;")


    def handle_radio_check(self):

        if not self.cloud_gemini_radio.isChecked():
            self.display_root.setText(CLIENT_PREFIX_TO_STRIP)
            self.display_root.setStyleSheet("color: black;")
            self.setWindowTitle(f"הדס לוי -  עורך דין - תוכנת חיפוש " + f" Hard Disk")
            self.non_cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_black)
            self.cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_gray)



        else:
            #white_cloud = '<span style="color: white;">☁️</span>'
            #black_bucket = '<span style="color: black;"> Bucket</span>'
            #self.display_root.setText(white_cloud + black_bucket)

            self.display_root.setText("☁️ Bucket")
            result = check_sync(CLIENT_PREFIX_TO_STRIP+"/גירושין/", BUCKET_NAME, prefix='גירושין')
            sync0 = result["sync!"]
            self.sync0 = sync0
            self.update_gcs_radio()
            if self.sync0:
                self.display_root.setStyleSheet("color: white; background-color: lightblue;")
            else:
                self.display_root.setStyleSheet("color: red; background-color: lightblue;")

            if self.update_app_title:
                self.set_window_title()

            self.cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_black)
            self.non_cloud_gemini_radio.setStyleSheet(CHECKBOX_STYLE_QSS_gray)

    def append_result(self, filepath, paragraph, search_terms):
        # Highlight search terms in paragraph
        highlighted_paragraph = highlight_terms(paragraph, search_terms)
        html_text = (
            # f"<p style='font-family:Arial; font-size:14pt;'>"
            f"<div dir='rtl' style='font-family:Arial; font-size:16pt;margin-right:40px;'>"
            f"<br>*<br>"
            # f"<b>קובץ:</b> {filepath}<br>"
            f"{highlighted_paragraph}"
            f"================================================================"
            # "<hr style='height:4px; border:none; background-color:red; margin:10px 0;'>"
            "</div>"

        )
        self.results_area.append(html_text)

    def append_result2(self, filepath, paragraph, search_terms, search_mode = ""):
        # Highlight search terms in paragraph
        if search_mode == 'chatgpt':

            html_text = (
                f"<div dir='rtl' style='font-family:Arial; font-size:16pt;margin-right:40px;'>"
                f"{paragraph}"
                "</div>"
            )
        else:
            highlighted_paragraph = highlight_terms(paragraph, search_terms)
            html_text = (
                f"<div dir='rtl' style='font-family:Arial; font-size:16pt;margin-right:40px;'>"
                f"{highlighted_paragraph}"
                "</div>"
            )
        self.results_area.append(html_text)


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = SearchApp()
    window.show()

    # The connection remains, pointing to the new, highly native handler.
    window.results_area.anchorClicked.connect(window._handle_link_click)

    sys.exit(app.exec_())